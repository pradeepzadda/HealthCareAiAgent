import streamlit as st
import json
import pandas as pd
from datetime import datetime
import os
import logging
import asyncio
import traceback
import time
import base64
from google.genai import types as genai_types
# Import settings early to configure authentication BEFORE importing agents
from settings import APP_NAME_FOR_ADK, USER_ID, INITIAL_STATE, ADK_SESSION_KEY, configure_adk_authentication, get_api_key
from database import get_database
from rag_service import get_rag_service

# Configure authentication early (this sets environment variables that ADK will use)
_ = configure_adk_authentication()

# Import your agents AFTER authentication is configured
from agent import (
    context_retriever_agent,
    test_case_generator_agent,
    compliance_agent,
    gap_analysis_agent,
    jira_formatter_agent,
    recommended_test_case_generator_agent,
)

#from services.adk_service import initialize_adk, run_adk_sync

# Import ADK components with proper error handling
try:
    from google.adk.runners import Runner
    from google.adk.sessions import InMemorySessionService
    ADK_AVAILABLE = True
    logging.info("Google ADK successfully imported")
except ImportError as e:
    logging.warning(f"Google ADK not available: {e}")
    ADK_AVAILABLE = False
    # Create dummy classes for type hints when ADK is not available
    class Runner:
        pass
    class InMemorySessionService:
        pass

# --- Logging Configuration ---
# Get the directory where app.py is located
_app_dir = os.path.dirname(os.path.abspath(__file__))
_log_dir = os.path.join(_app_dir, 'logs')
os.makedirs(_log_dir, exist_ok=True)  # Create logs directory

_main_log_file_path = os.path.join(_app_dir, 'healthcare_agent.log')

# Session-specific log handlers dictionary
_session_log_handlers = {}

class SessionContextFilter(logging.Filter):
    """Filter to add session ID to log records."""
    def filter(self, record):
        # Try to get session ID from context or thread-local storage
        if hasattr(record, 'session_id'):
            return True
        # Try to get from logging context if available
        record.session_id = getattr(logging.getLogger(), 'session_id', 'N/A')
        return True

# Configure logging with file handler
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - [%(session_id)s] - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(_main_log_file_path, mode='a', encoding='utf-8')
    ],
    force=True  # Force reconfiguration if logging was already configured
)

# Add session context filter to root logger
for handler in logging.root.handlers:
    handler.addFilter(SessionContextFilter())

logger = logging.getLogger(__name__)
logger.info(f"Logging initialized. Main log file: {_main_log_file_path}, Session logs directory: {_log_dir}")

def get_session_log_file(session_id: str) -> str:
    """Get the log file path for a specific session."""
    # Clean session ID for filename (remove special characters)
    safe_session_id = session_id.replace('/', '_').replace('\\', '_')
    return os.path.join(_log_dir, f'session_{safe_session_id}.log')

def add_session_log_handler(session_id: str):
    """Add a session-specific log file handler."""
    global _session_log_handlers
    
    if session_id in _session_log_handlers:
        return  # Already added
    
    try:
        session_log_file = get_session_log_file(session_id)
        handler = logging.FileHandler(session_log_file, mode='a', encoding='utf-8')
        
        # Use format with session ID
        formatter = logging.Formatter(
            '%(asctime)s - [%(session_id)s] - %(name)s - %(levelname)s - %(message)s'
        )
        handler.setFormatter(formatter)
        handler.addFilter(SessionContextFilter())
        
        # Add to root logger
        logging.root.addHandler(handler)
        _session_log_handlers[session_id] = handler
        
        logger.info(f"Session log handler added. Log file: {session_log_file}")
    except Exception as e:
        logger.warning(f"Could not add session log handler for {session_id}: {e}")

def remove_session_log_handler(session_id: str):
    """Remove a session-specific log file handler."""
    global _session_log_handlers
    
    if session_id in _session_log_handlers:
        try:
            handler = _session_log_handlers[session_id]
            logging.root.removeHandler(handler)
            handler.close()
            del _session_log_handlers[session_id]
            logger.info(f"Session log handler removed for {session_id}")
        except Exception as e:
            logger.warning(f"Could not remove session log handler for {session_id}: {e}")

def set_logging_session_id(session_id: str):
    """Set the session ID in logging context for this thread."""
    # Store in logger instance for filter to access
    logging.getLogger().session_id = session_id

#CSS components
def show_loading_animation(title="Processing", subtitle="Please wait...", description="", animation_type="dots"):
    """Show custom loading animation in center of page"""
    
    # Choose animation based on type
    if animation_type == "dots":
        animation_html = '''
        <div class="dots-spinner">
            <div></div>
            <div></div>
            <div></div>
            <div></div>
        </div>
        '''
    elif animation_type == "ring":
        animation_html = '''
        <div class="ring-spinner">
            <div></div>
            <div></div>
            <div></div>
            <div></div>
        </div>
        '''
    elif animation_type == "dna":
        animation_html = '''
        <div class="dna-spinner">
            <div></div>
            <div></div>
        </div>
        '''
    else:
        animation_html = '''
        <div class="dots-spinner">
            <div></div>
            <div></div>
            <div></div>
            <div></div>
        </div>
        '''
    
    loading_css_js = f'''
    <style>
        .loading-overlay {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(255, 255, 255, 0.95);
            z-index: 9999;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        }}

        .loading-content {{
            text-align: center;
            max-width: 400px;
            padding: 2rem;
        }}

        .loading-animation {{
            width: 80px;
            height: 80px;
            margin: 0 auto 2rem;
        }}

        /* Spinning dots animation */
        .dots-spinner {{
            display: inline-block;
            position: relative;
            width: 80px;
            height: 80px;
        }}

        .dots-spinner div {{
            position: absolute;
            top: 33px;
            width: 13px;
            height: 13px;
            border-radius: 50%;
            background: #ff6b6b;
            animation-timing-function: cubic-bezier(0, 1, 1, 0);
        }}

        .dots-spinner div:nth-child(1) {{
            left: 8px;
            animation: dots1 0.6s infinite;
        }}

        .dots-spinner div:nth-child(2) {{
            left: 8px;
            animation: dots2 0.6s infinite;
        }}

        .dots-spinner div:nth-child(3) {{
            left: 32px;
            animation: dots2 0.6s infinite;
        }}

        .dots-spinner div:nth-child(4) {{
            left: 56px;
            animation: dots3 0.6s infinite;
        }}

        @keyframes dots1 {{
            0% {{ transform: scale(0); }}
            100% {{ transform: scale(1); }}
        }}

        @keyframes dots3 {{
            0% {{ transform: scale(1); }}
            100% {{ transform: scale(0); }}
        }}

        @keyframes dots2 {{
            0% {{ transform: translate(0, 0); }}
            100% {{ transform: translate(24px, 0); }}
        }}

        /* Pulsing ring animation */
        .ring-spinner {{
            display: inline-block;
            position: relative;
            width: 80px;
            height: 80px;
        }}

        .ring-spinner div {{
            box-sizing: border-box;
            display: block;
            position: absolute;
            width: 64px;
            height: 64px;
            margin: 8px;
            border: 8px solid #4ecdc4;
            border-radius: 50%;
            animation: ring-spin 1.2s cubic-bezier(0.5, 0, 0.5, 1) infinite;
            border-color: #4ecdc4 transparent transparent transparent;
        }}

        .ring-spinner div:nth-child(1) {{ animation-delay: -0.45s; }}
        .ring-spinner div:nth-child(2) {{ animation-delay: -0.3s; }}
        .ring-spinner div:nth-child(3) {{ animation-delay: -0.15s; }}

        @keyframes ring-spin {{
            0% {{ transform: rotate(0deg); }}
            100% {{ transform: rotate(360deg); }}
        }}

        /* DNA helix animation */
        .dna-spinner {{
            display: inline-block;
            position: relative;
            width: 80px;
            height: 80px;
        }}

        .dna-spinner div {{
            position: absolute;
            width: 6px;
            height: 6px;
            border-radius: 50%;
            background: #45b7d1;
        }}

        .dna-spinner div:nth-child(1) {{ top: 10px; left: 10px; animation: dna1 2s linear infinite; }}
        .dna-spinner div:nth-child(2) {{ top: 10px; right: 10px; animation: dna2 2s linear infinite; }}

        @keyframes dna1 {{
            25% {{ top: 10px; left: 64px; }}
            50% {{ top: 64px; left: 64px; }}
            75% {{ top: 64px; left: 10px; }}
            100% {{ top: 10px; left: 10px; }}
        }}

        @keyframes dna2 {{
            25% {{ top: 64px; right: 64px; }}
            50% {{ top: 64px; right: 10px; }}
            75% {{ top: 10px; right: 10px; }}
            100% {{ top: 10px; right: 64px; }}
        }}

        .loading-title {{
            font-size: 1.8rem;
            font-weight: 600;
            color: #2c3e50;
            margin-bottom: 1rem;
            background: linear-gradient(45deg, #667eea 0%, #764ba2 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }}

        .loading-subtitle {{
            font-size: 1.1rem;
            color: #7f8c8d;
            margin-bottom: 0.5rem;
            font-weight: 400;
        }}

        .loading-description {{
            font-size: 0.9rem;
            color: #95a5a6;
            line-height: 1.4;
        }}

        .progress-bar {{
            width: 100%;
            height: 4px;
            background: #ecf0f1;
            border-radius: 2px;
            margin-top: 1.5rem;
            overflow: hidden;
        }}

        .progress-fill {{
            height: 100%;
            background: linear-gradient(90deg, #667eea, #764ba2);
            border-radius: 2px;
            animation: progress 3s ease-in-out infinite;
        }}

        @keyframes progress {{
            0% {{ width: 0%; }}
            50% {{ width: 70%; }}
            100% {{ width: 100%; }}
        }}

        .pulse-text {{
            animation: pulse 2s ease-in-out infinite;
        }}

        @keyframes pulse {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.6; }}
        }}
    </style>
    
    <div class="loading-overlay" id="loading-overlay">
        <div class="loading-content">
            <div class="loading-animation">
                {animation_html}
            </div>
            <h2 class="loading-title">{title}</h2>
            <p class="loading-subtitle pulse-text">{subtitle}</p>
            <p class="loading-description">{description}</p>
            <div class="progress-bar">
                <div class="progress-fill"></div>
            </div>
        </div>
    </div>
    
    <script>
        // Auto-hide after 30 seconds as fallback
        setTimeout(function() {{
            var overlay = document.getElementById('loading-overlay');
            if (overlay) {{
                overlay.style.display = 'none';
            }}
        }}, 30000);
    </script>
    '''
    
    return st.markdown(loading_css_js, unsafe_allow_html=True)

def hide_loading_animation():
    """Hide the loading animation"""
    hide_css = '''
    <style>
        .loading-overlay {
            display: none !important;
        }
    </style>
    '''
    return st.markdown(hide_css, unsafe_allow_html=True)

# Context manager for easy use
class LoadingAnimation:
    def __init__(self, title="Processing", subtitle="Please wait...", description="", animation_type="dots"):
        self.title = title
        self.subtitle = subtitle
        self.description = description
        self.animation_type = animation_type
        self.placeholder = None
    
    def __enter__(self):
        # Create a placeholder that will be replaced
        self.placeholder = st.empty()
        with self.placeholder:
            show_loading_animation(
                title=self.title,
                subtitle=self.subtitle, 
                description=self.description,
                animation_type=self.animation_type
            )
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        # Clear the placeholder
        self.placeholder.empty()

def show_loading_overlay(message="Processing...", submessage="Please wait"):
    """Simple loading overlay that works reliably in Streamlit"""
    st.markdown(f"""
    <div style='
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background-color: rgba(255, 255, 255, 0.9);
        z-index: 999;
        display: flex;
        justify-content: center;
        align-items: center;
        flex-direction: column;
    '>
        <div style='text-align: center;'>
            <div style='
                border: 4px solid #f3f3f3;
                border-top: 4px solid #3498db;
                border-radius: 50%;
                width: 60px;
                height: 60px;
                animation: spin 2s linear infinite;
                margin: 0 auto 20px;
            '></div>
            <h3 style='color: #2c3e50; margin: 10px 0;'>{message}</h3>
            <p style='color: #7f8c8d;'>{submessage}</p>
        </div>
    </div>
    <style>
    @keyframes spin {{
        0% {{ transform: rotate(0deg); }}
        100% {{ transform: rotate(360deg); }}
    }}
    </style>
    """, unsafe_allow_html=True)

# Context manager for easy use
class SimpleLoadingOverlay:
    def __init__(self, message="Processing...", submessage="Please wait"):
        self.message = message
        self.submessage = submessage
        self.container = None
    
    def __enter__(self):
        self.container = st.empty()
        with self.container:
            show_loading_overlay(self.message, self.submessage)
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.container:
            self.container.empty()
# --- ADK Runtime Configuration ---
MAX_RETRIES = 3
RETRY_DELAY = 1

@st.cache_resource
def create_adk_session(session_id):
    """
    Initializes the Google ADK Runner and manages the ADK session.
    Uses Streamlit's cache_resource to ensure this runs only once per app load.
    """
    try:
        asyncio.run(
            session_service.create_session(
                app_name=APP_NAME_FOR_ADK,
                user_id=USER_ID,
                session_id=session_id,
                state=INITIAL_STATE, # Initialize with predefined state.
            )
        )
        return True
    except Exception as e:
        logger.error(f"Exception creating ADK session: {e}")
        return False

@st.cache_resource
def get_adk_session(session_id):
    try:
        # If an ADK session ID already exists (e.g., on a Streamlit rerun), retrieve it.
        # Verify if the session still exists in the ADK session service.
        # This handles cases where the service might reset (less common with cache_resource, but good practice).
        if not asyncio.run(session_service.get_session(app_name=APP_NAME_FOR_ADK, user_id=USER_ID, session_id=session_id)):
            # If the session was lost (e.g., full app restart without clearing cache), recreate it.
            asyncio.run(
                session_service.create_session(
                    app_name=APP_NAME_FOR_ADK,
                    user_id=USER_ID,
                    session_id=session_id,
                    state=INITIAL_STATE
                )
            )
        return True
    except Exception as e:
        logger.error(f"Exception creating ADK session: {e}")
        return False


async def run_adk_async(runner: Runner, session_id: str, user_message_text: str):
    """
    Asynchronously runs a single turn of the ADK agent conversation.
    """
    # Set session ID in logging context for this operation
    set_logging_session_id(session_id)
    
    try:
        session = await runner.session_service.get_session(app_name=APP_NAME_FOR_ADK, user_id=USER_ID, session_id=session_id)
        if not session:
            return "Error: ADK session not found."
        
        # Prepare the user's message in the format expected by ADK/Gemini.
        content = genai_types.Content(role='user', parts=[genai_types.Part(text=user_message_text)])
        final_response_text = None
        all_responses = []
        event_count = 0
        
        # Iterate through the asynchronous events generated by the ADK runner.
        # ADK can yield multiple events (e.g., tool calls, interim responses) before the final response.
        try:
            async for event in runner.run_async(user_id=USER_ID, session_id=session_id, new_message=content):
                event_count += 1
                
                # Log event type for debugging
                event_type = type(event).__name__
                is_final = hasattr(event, 'is_final_response') and event.is_final_response()
                logger.info(f"Event #{event_count}: {event_type}, is_final_response: {is_final}")
                
                # Try to extract text from ANY event that might have content
                # (even if not final, as some responses might not be marked as final)
                event_text = None
                
                # Method 1: event.content.parts[].text
                if hasattr(event, 'content') and event.content:
                    if hasattr(event.content, 'parts') and event.content.parts:
                        for part in event.content.parts:
                            if hasattr(part, 'text') and part.text:
                                event_text = part.text
                                logger.info(f"Extracted text from event.content.parts[].text: {len(event_text)} chars")
                                break
                    elif hasattr(event.content, 'text') and event.content.text:
                        event_text = event.content.text
                        logger.info(f"Extracted text from event.content.text: {len(event_text)} chars")
                
                # Method 2: event.text directly
                if not event_text and hasattr(event, 'text') and event.text:
                    event_text = event.text
                    logger.info(f"Extracted text from event.text: {len(event_text)} chars")
                
                # Method 3: Try to get text from event directly via str() or other attributes
                if not event_text:
                    # Try common attributes
                    for attr in ['message', 'response', 'output', 'result', 'data']:
                        if hasattr(event, attr):
                            attr_value = getattr(event, attr)
                            if isinstance(attr_value, str) and attr_value.strip():
                                event_text = attr_value
                                logger.info(f"Extracted text from event.{attr}: {len(event_text)} chars")
                                break
                
                # Method 4: Check for errors or special finish reasons
                if not event_text and is_final:
                    # Check if there's an error or unexpected finish reason
                    if hasattr(event, 'finish_reason'):
                        finish_reason = event.finish_reason
                        if finish_reason:
                            finish_str = str(finish_reason)
                            logger.warning(f"Event has finish_reason: {finish_str}")
                            
                            # Check for error conditions
                            if 'UNEXPECTED_TOOL_CALL' in finish_str or 'ERROR' in finish_str:
                                error_msg = f"Agent encountered an error: {finish_str}"
                                if hasattr(event, 'error_message') and event.error_message:
                                    error_msg += f" - {event.error_message}"
                                if hasattr(event, 'error_code') and event.error_code:
                                    error_msg += f" (code: {event.error_code})"
                                logger.error(error_msg)
                                # Don't extract this as valid content
                                event_text = None
                            else:
                                logger.info(f"Finish reason (non-error): {finish_str}")
                    
                    # Check if content is explicitly None
                    if hasattr(event, 'content'):
                        if event.content is None:
                            logger.warning("Event content is None - this indicates an error or empty response")
                            # Try to find error information
                            error_info = []
                            if hasattr(event, 'finish_reason'):
                                error_info.append(f"finish_reason: {event.finish_reason}")
                            if hasattr(event, 'error_code'):
                                error_info.append(f"error_code: {event.error_code}")
                            if hasattr(event, 'error_message'):
                                error_info.append(f"error_message: {event.error_message}")
                            if error_info:
                                event_text = f"Error: {'; '.join(error_info)}"
                                logger.error(f"Extracted error info: {event_text}")
                    
                    # Last resort - try converting event to string (but be very careful)
                    if not event_text and is_final:
                        try:
                            event_str = str(event)
                            if event_str and event_str.strip() and len(event_str) > 50:
                                # Only use if it looks like actual content, not just object representation
                                # Check for signs it's an object representation, not content
                                if (not event_str.startswith('<') and 
                                    'object at 0x' not in event_str and
                                    'content=None' not in event_str and
                                    'FinishReason' not in event_str and
                                    'UNEXPECTED_TOOL_CALL' not in event_str):
                                    event_text = event_str
                                    logger.info(f"Extracted text from str(event): {len(event_text)} chars")
                                else:
                                    # This is an error event, not content
                                    logger.error(f"Rejected str(event) - appears to be error/object representation")
                                    logger.error(f"Preview: {event_str[:300]}")
                                    
                                    # Build proper error message
                                    error_parts = ["Agent returned an error event instead of content:"]
                                    if 'UNEXPECTED_TOOL_CALL' in event_str:
                                        error_parts.append("UNEXPECTED_TOOL_CALL - The agent tried to use a tool but encountered an error.")
                                    if 'content=None' in event_str:
                                        error_parts.append("No content was generated.")
                                    if hasattr(event, 'error_message') and event.error_message:
                                        error_parts.append(f"Error message: {event.error_message}")
                                    
                                    event_text = "Error: " + " ".join(error_parts)
                                    logger.error(f"Error message constructed: {event_text}")
                        except Exception as str_error:
                            logger.error(f"Error converting event to string: {str_error}")
                            pass
                
                # Store the text if we found it
                if event_text:
                    if is_final:
                        final_response_text = event_text
                        logger.info(f"Final response captured from event #{event_count}: {len(event_text)} chars")
                        break
                    else:
                        all_responses.append(event_text)
                        logger.debug(f"Non-final response part: {len(event_text)} chars")
                elif is_final:
                    # Final response but couldn't extract text - log the event structure in detail
                    logger.error(f"⚠️ Final response event #{event_count} but could not extract text! Event type: {event_type}")
                    
                    # Log all event attributes
                    attrs = [attr for attr in dir(event) if not attr.startswith('_')]
                    logger.error(f"Event attributes: {attrs}")
                    
                    # Try to inspect the content structure in detail
                    if hasattr(event, 'content'):
                        logger.error(f"event.content exists: {event.content}")
                        logger.error(f"event.content type: {type(event.content)}")
                        if hasattr(event.content, 'parts'):
                            logger.error(f"event.content.parts: {event.content.parts}")
                            logger.error(f"event.content.parts type: {type(event.content.parts)}")
                            if event.content.parts:
                                for i, part in enumerate(event.content.parts):
                                    logger.error(f"Part {i} type: {type(part)}, attrs: {[a for a in dir(part) if not a.startswith('_')]}")
                                    if hasattr(part, 'text'):
                                        logger.error(f"Part {i} has text attr: {part.text}")
                                    else:
                                        logger.error(f"Part {i} no text attr")
                    
                    # Try repr
                    try:
                        event_repr = repr(event)
                        logger.error(f"Event repr (first 1000 chars): {event_repr[:1000]}")
                    except Exception as repr_error:
                        logger.error(f"Could not get event repr: {repr_error}")
                    
                    # Try to get the full event structure as JSON if possible
                    try:
                        import json
                        if hasattr(event, '__dict__'):
                            logger.error(f"Event __dict__: {json.dumps({k: str(v)[:200] for k, v in event.__dict__.items()}, indent=2)}")
                    except:
                        pass
        
        except Exception as async_error:
            logger.error(f"Error in async event loop: {async_error}", exc_info=True)
            # If we got some responses before the error, try to use the last one
            if all_responses:
                final_response_text = all_responses[-1]
                logger.warning(f"Using last collected response after async error: {len(final_response_text)} chars")
            else:
                raise
        
        # Log what we collected
        logger.info(f"Event loop completed. Total events: {event_count}, Final response captured: {final_response_text is not None}")
        
        # If no events were received at all, log a warning
        if event_count == 0:
            error_msg = "No events received from ADK runner. The agent may not have started or the async generator is not yielding events."
            logger.error(error_msg)
            logger.error(f"Session ID: {session_id}, Runner: {runner}, Message length: {len(user_message_text)}")
            # Force flush logs
            for handler in logger.handlers:
                handler.flush()
            return f"Error: {error_msg}"
        
        # If no final response, check if we collected any responses
        if not final_response_text:
            if all_responses:
                final_response_text = "\n".join(all_responses)
                logger.warning(f"No final response found, using all collected responses: {len(final_response_text)} chars")
            else:
                error_msg = f"No response captured after {event_count} events. This likely indicates a tool call error (UNEXPECTED_TOOL_CALL) or agent failure."
                logger.error(error_msg)
                logger.error(f"Event count: {event_count}, Session ID: {session_id}")
                logger.error(f"All collected responses: {all_responses}")
                logger.error("This often happens when an agent tries to use a tool (like RAG search) but encounters an error.")
                # Force flush logs
                for handler in logger.handlers:
                    handler.flush()
                return f"Error: Agent encountered an UNEXPECTED_TOOL_CALL error. The agent tried to use a tool (possibly RAG) but failed. Check logs for details. Received {event_count} events but no valid content."
        
        # Validate the response
        if not final_response_text.strip():
            logger.error("Agent returned empty response text")
            return "Error: Agent returned an empty response. Please check logs for details."
        
        logger.info(f"Successfully captured agent response: {len(final_response_text)} characters")
        return final_response_text
        
    except Exception as e:
        error_msg = f"Error running agent: {str(e)}"
        logger.error(error_msg, exc_info=True)
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"Session ID: {session_id}, Message preview: {user_message_text[:100]}")
        # Force flush logs to ensure they're written
        for handler in logger.handlers:
            handler.flush()
        return f"Error: {error_msg}"

def run_adk_sync(runner: Runner, session_id: str, user_message_text: str) -> str:
    """
    Synchronous wrapper for running ADK, as Streamlit does not directly support async calls in the main thread.
    """
    # Runs the asynchronous function in a new event loop.
    return asyncio.run(run_adk_async(runner, session_id, user_message_text))

@st.cache_resource
def get_session_service():
    """Initializes and returns a cached InMemorySessionService instance."""
    if not ADK_AVAILABLE:
        logger.warning("ADK not available, returning None for session service")
        return None
    try:
        service = InMemorySessionService()
        logger.info("InMemorySessionService initialized successfully")
        return service
    except Exception as e:
        logger.error(f"Failed to initialize Session Service: {e}")
        return None

# Initialize components with better error handling
session_service = get_session_service()
runners_initialized = False

def initialize_runners():
    """Initialize runners with proper error handling and retries."""
    global runners_initialized
    global context_retriever_runner, test_case_generator_runner, compliance_runner, gap_analysis_runner, jira_formatter_runner, recommended_test_case_generator_runner
    
    if not ADK_AVAILABLE or not session_service:
        logger.warning("Cannot initialize runners: ADK or session service not available")
        return False
    
    if runners_initialized:
        return True
    
    for attempt in range(MAX_RETRIES):
        try:
            logger.info(f"Initializing runners, attempt {attempt + 1}")
            
            context_retriever_runner = Runner(
                app_name=APP_NAME_FOR_ADK, 
                agent=context_retriever_agent, 
                session_service=session_service
            )
            test_case_generator_runner = Runner(
                app_name=APP_NAME_FOR_ADK, 
                agent=test_case_generator_agent, 
                session_service=session_service
            )
            compliance_runner = Runner(
                app_name=APP_NAME_FOR_ADK, 
                agent=compliance_agent, 
                session_service=session_service
            )
            gap_analysis_runner = Runner(
                app_name=APP_NAME_FOR_ADK, 
                agent=gap_analysis_agent, 
                session_service=session_service
            )
            jira_formatter_runner = Runner(
                app_name=APP_NAME_FOR_ADK, 
                agent=jira_formatter_agent, 
                session_service=session_service
            )
            recommended_test_case_generator_runner = Runner(
                app_name=APP_NAME_FOR_ADK, 
                agent=recommended_test_case_generator_agent, 
                session_service=session_service
            )
            
            runners_initialized = True
            logger.info("All runners initialized successfully")
            return True
            
        except Exception as e:
            logger.error(f"Failed to initialize runners (attempt {attempt + 1}): {e}")
            if attempt < MAX_RETRIES - 1:
                time.sleep(RETRY_DELAY)
            else:
                logger.error("All attempts to initialize runners failed")
                return False
    
    return False

# --- Helper Functions for Document Parsing ---
def parse_pdf(file):
    """Parse PDF file and extract text with better error handling."""
    try:
        # Reset file pointer to beginning
        file.seek(0)
        import pypdf
        logger.info(f"Parsing PDF file: {file.name}")
        pdf_reader = pypdf.PdfReader(file)
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
                logger.debug(f"Extracted text from page {page_num}: {len(page_text)} characters")
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")
                continue
        
        final_text = "\n".join(text_parts).strip()
        logger.info(f"PDF parsing completed: {len(final_text)} total characters")
        return final_text if final_text else None
        
    except ImportError:
        error_msg = "pypdf package is required for PDF parsing. Install with: pip install pypdf"
        logger.error(error_msg)
        st.error(error_msg)
        return None
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        st.error(f"Error parsing PDF: {e}")
        return None

def parse_docx(file):
    """Parse DOCX file and extract text with better error handling."""
    try:
        # Reset file pointer to beginning
        file.seek(0)
        import docx
        logger.info(f"Parsing DOCX file: {file.name}")
        doc = docx.Document(file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        final_text = "\n".join(text_parts)
        logger.info(f"DOCX parsing completed: {len(final_text)} characters")
        return final_text if final_text else None
        
    except ImportError:
        error_msg = "python-docx package is required for DOCX parsing. Install with: pip install python-docx"
        logger.error(error_msg)
        st.error(error_msg)
        return None
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        st.error(f"Error parsing DOCX: {e}")
        return None

def parse_txt(file):
    """Parse TXT file and extract text with encoding detection."""
    try:
        # Reset file pointer to beginning
        file.seek(0)
        logger.info(f"Parsing TXT file: {file.name}")
        
        # Read as bytes first
        content_bytes = file.read()
        
        # Try UTF-8 first
        try:
            content = content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            # Fallback to other encodings
            try:
                content = content_bytes.decode("latin-1")
                logger.info("Used latin-1 encoding for text file")
            except UnicodeDecodeError:
                content = content_bytes.decode("utf-8", errors="ignore")
                logger.warning("Used UTF-8 with error ignoring for text file")
        
        logger.info(f"TXT parsing completed: {len(content)} characters")
        return content.strip() if content.strip() else None
        
    except Exception as e:
        logger.error(f"Error parsing TXT: {e}")
        st.error(f"Error parsing TXT: {e}")
        return None

# --- RAG Compliance Document Upload Handler ---
def handle_compliance_document_upload(uploaded_file):
    """Handle compliance document upload and ingestion into RAG."""
    try:
        # Parse file
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        parsers = {
            ".pdf": parse_pdf,
            ".docx": parse_docx,
            ".txt": parse_txt
        }
        
        parser = parsers.get(file_extension, parse_txt)
        file_content = parser(uploaded_file)
        
        if not file_content:
            st.error("Could not extract meaningful content from the compliance document.")
            return False
        
        # Get RAG service and ingest document
        rag_service = get_rag_service()
        
        if not rag_service:
            st.error("RAG service is not available. Please check installation of chromadb and sentence-transformers.")
            return False
        
        # Ingest into RAG
        with st.spinner(f"Ingesting {uploaded_file.name} into RAG system..."):
            num_chunks = rag_service.ingest_document(
                document_text=file_content,
                document_name=uploaded_file.name,
                metadata={
                    "upload_date": datetime.now().isoformat(),
                    "file_type": file_extension
                }
            )
        
        st.success(f"Successfully ingested '{uploaded_file.name}' into RAG system ({num_chunks} chunks created)")
        return True
        
    except Exception as e:
        error_msg = f"Failed to ingest compliance document: {e}"
        logger.error(error_msg)
        st.error(error_msg)
        return False

# --- Database Initialization (SQLite) ---
@st.cache_resource
def get_db_client():
    """Initialize SQLite database client with better error handling."""
    try:
        db = get_database()
        if db and db.is_available():
            logger.info("SQLite database initialized and tested successfully")
            return db
        else:
            logger.warning("SQLite database not available")
            st.warning("Database is not available. Session data will not be persisted.")
            return None
    except Exception as e:
        logger.warning(f"Failed to initialize database: {e}")
        st.warning("Database is not available. Session data will not be persisted.")
        return None

db = get_db_client()

# --- Async Helper Functions ---
async def create_session_async(app_name, user_id, session_id):
    """Async session creation with fallback."""
    try:
        await session_service.create_session(
            app_name=app_name,
            user_id=user_id,
            session_id=session_id
        )
        return True
    except Exception as e:
        logger.warning(f"Session creation with app_name failed: {e}")
        # Try without app_name
        try:
            await session_service.create_session(
                user_id=user_id,
                session_id=session_id
            )
            return True
        except Exception as e2:
            logger.error(f"Session creation without app_name also failed: {e2}")
            return False

async def get_session_async(app_name, user_id, session_id):
    """Async session retrieval with fallback."""
    try:
        return await session_service.get_session(
            app_name=app_name,
            user_id=user_id,
            session_id=session_id
        )
    except Exception as e:
        logger.warning(f"Session retrieval failed: {e}")
        return None

def run_async_function(coro):
    """Run async function in sync context with proper event loop handling."""
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # We're in an async context, create a new event loop in a thread
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, coro)
                return future.result(timeout=30)  # 30 second timeout
        else:
            return loop.run_until_complete(coro)
    except RuntimeError:
        # No event loop running, create a new one
        return asyncio.run(coro)
    except Exception as e:
        logger.error(f"Error running async function: {e}")
        raise

# --- Session State Management ---
def initialize_session_state():
    """Initialize session state with proper error handling."""
    # Initialize ADK session ID if it doesn't exist
    if ADK_SESSION_KEY not in st.session_state:
        session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S%f')[:-3]}"
        st.session_state[ADK_SESSION_KEY] = session_id
        
        # Set session ID in logging context
        set_logging_session_id(session_id)
        
        # Add session-specific log handler
        add_session_log_handler(session_id)
        
        if create_adk_session(session_id):
            logger.info(f"Created new ADK session: {session_id}")
        else:
            logger.error(f"Failed to create new ADK session: {session_id}")
    else:
        # On a rerun, just ensure the ADK session is still valid.
        session_id = st.session_state[ADK_SESSION_KEY]
        set_logging_session_id(session_id)
        add_session_log_handler(session_id)
        get_adk_session(session_id)

    # Initialize app-specific state variables if they don't exist.
    defaults = {
        'context': None,
        'test_cases': [],
        'gap_analysis': "",
        'jira_csv': "",
        'processing_complete': False,
        'last_error': None,
        'requirement_contexts': {},  # Store per-requirement contexts
        'recommended_test_cases': [],  # Recommended test cases from gap analysis
        'included_recommendations': {},  # Dict mapping idx -> True/False for include/exclude
        'edited_recommendations': {},  # Edited versions of recommendations (all are editable)
        'edited_test_cases': {},  # Dict mapping idx -> edited test case data (for generated test cases)
        'included_test_cases': {},  # Dict mapping idx -> True/False for include/exclude (for generated test cases)
        'finalized_test_cases': [],  # Final list of test cases after editing and filtering
        'is_finalized': False,  # Flag to indicate if test cases have been finalized
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

def reset_run_state():
    """Reset state for new run with better session management."""
    # Clear all app-specific state variables
    for key in ['context', 'test_cases', 'gap_analysis', 'jira_csv', 'processing_complete', 'last_error', 
                'requirement_contexts', 'recommended_test_cases', 'included_recommendations', 
                'edited_recommendations', 'edited_test_cases', 'included_test_cases', 
                'finalized_test_cases', 'is_finalized']:
        if key in st.session_state:
            del st.session_state[key]
    
    # Create a new ADK session ID for the new run
    if ADK_AVAILABLE and session_service:

        old_session_id = st.session_state.get(ADK_SESSION_KEY)
        new_session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S%f')[:-3]}"
        st.session_state[ADK_SESSION_KEY] = new_session_id
        
        try:
            # .clear() is needed for @st.cache_resource to re-run the function
            create_adk_session.clear()
            success = create_adk_session(new_session_id)
            if success:
                logger.info(f"Created new ADK session: {new_session_id}")
            else:
                logger.error("Failed to create new session, reverting to old session")
                if old_session_id:
                    st.session_state[ADK_SESSION_KEY] = old_session_id
        except Exception as e:
            logger.error(f"Exception creating new session: {e}")
            if old_session_id:
                st.session_state[ADK_SESSION_KEY] = old_session_id

def run_agent_with_retry(runner, prompt_text: str, max_retries: int = 3) -> str:
    """Run agent with retry logic and better error handling."""
    if not ADK_AVAILABLE:
        raise RuntimeError("Google ADK is not available. Cannot run agents.")
    
    if not runners_initialized and not initialize_runners():
        raise RuntimeError("Failed to initialize runners.")
    
    last_exception = None
    
    for attempt in range(max_retries):
        try:
            logger.info(f"Running agent, attempt {attempt + 1}")
            
            # Ensure session exists
            # session = run_async_function(get_session_async(
            #     app_name=APP_NAME,
            #     user_id=st.session_state.user_id,
            #     session_id=st.session_state.session_id
            # ))
            
            # if session is None:
            #     success = run_async_function(create_session_async(
            #         app_name=APP_NAME,
            #         user_id=st.session_state.user_id,
            #         session_id=st.session_state.session_id
            #     ))
            #     if not success:
            #         raise RuntimeError("Could not create or find session")
            
            # Try different message formats
            message_formats = [
                prompt_text,  # Simple string
                {"text": prompt_text},  # Text wrapper
                {"content": prompt_text},  # Content wrapper
            ]
            
            for i, msg_format in enumerate(message_formats):
                try:
                    logger.debug(f"Trying message format {i+1}")
                    result = runner.run(
                        user_id=USER_ID,
                        session_id=st.session_state[ADK_SESSION_KEY],
                        new_message=msg_format
                    )
                    
                    # Process result
                    if hasattr(result, '__iter__') and not isinstance(result, (str, bytes)):
                        response_parts = []
                        for chunk in result:
                            if isinstance(chunk, str):
                                response_parts.append(chunk)
                            elif hasattr(chunk, 'text'):
                                response_parts.append(chunk.text)
                            elif hasattr(chunk, 'content'):
                                response_parts.append(chunk.content)
                            else:
                                chunk_str = str(chunk)
                                if chunk_str and chunk_str != str(type(chunk)):
                                    response_parts.append(chunk_str)
                        
                        final_response = ''.join(response_parts)
                        if final_response.strip():
                            logger.info(f"Successfully collected response: {len(final_response)} characters")
                            return final_response
                    elif isinstance(result, str) and result.strip():
                        return result
                    elif hasattr(result, 'text') and result.text.strip():
                        return result.text
                    elif hasattr(result, 'content') and result.content.strip():
                        return result.content
                        
                except Exception as e:
                    logger.warning(f"Message format {i+1} failed: {e}")
                    if i == len(message_formats) - 1:  # Last format
                        raise e
                    continue
                    
            raise RuntimeError("All message formats failed")
            
        except Exception as e:
            last_exception = e
            logger.warning(f"Agent run attempt {attempt + 1} failed: {e}")
            
            if attempt < max_retries - 1:
                logger.info(f"Retrying in {RETRY_DELAY} seconds...")
                time.sleep(RETRY_DELAY)
            else:
                logger.error("All retry attempts exhausted")
    
    raise RuntimeError(f"Agent execution failed after {max_retries} attempts. Last error: {last_exception}")

def parse_recommendations_from_gap_analysis(gap_analysis_text: str) -> list:
    """
    Parse recommended test cases from gap analysis markdown.
    Returns a list of dictionaries with test case structure.
    Only extracts complete, structured test cases (with Test Case ID, Summary, Test Steps, Expected Result).
    """
    import re
    
    recommendations = []
    
    if not gap_analysis_text or not gap_analysis_text.strip():
        return recommendations
    
    try:
        # Find the "Recommendations for New Test Cases" section
        rec_section_pattern = r'(?:^|\n)#{1,4}\s*Recommendations?\s+for\s+New\s+Test\s+Cases?[\s\S]*?(?=\n#{1,4}\s|$)'
        rec_section_match = re.search(rec_section_pattern, gap_analysis_text, re.IGNORECASE | re.MULTILINE)
        
        if not rec_section_match:
            # Try alternative pattern without heading markers
            rec_section_pattern = r'Recommendations?\s+for\s+New\s+Test\s+Cases?[\s\S]*?(?=\n\n#{1,4}\s|$)'
            rec_section_match = re.search(rec_section_pattern, gap_analysis_text, re.IGNORECASE)
        
        if not rec_section_match:
            logger.warning("Could not find 'Recommendations for New Test Cases' section")
            return recommendations
        
        rec_section = rec_section_match.group(0)
        
        # First, try to parse markdown tables (new format)
        # Look for markdown table with columns: Test Case Summary, Test Steps, Expected Result, Priority
        lines = rec_section.split('\n')
        
        # Track current requirement ID as we iterate through lines
        current_req_id = None
        in_table = False
        header_found = False
        separator_passed = False
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Look for gap headers to extract requirement ID
            gap_header_match = re.search(r'###\s+For\s+Gap\s+\d+[:\s]+.*?(REQ-?\d+)', line_stripped, re.IGNORECASE)
            if gap_header_match:
                current_req_id = gap_header_match.group(1)
                in_table = False
                header_found = False
                separator_passed = False
                continue
            
            # Check if this is a table header row
            if re.search(r'Test\s+Case\s+Summary.*?Test\s+Steps.*?Expected\s+Result.*?Priority', line_stripped, re.IGNORECASE):
                header_found = True
                in_table = True
                separator_passed = False
                continue
            
            # Check if this is a separator row (contains dashes and pipes)
            if header_found and re.search(r'^\|[\s\-:]+\|', line_stripped) and re.search(r'[-:]', line_stripped):
                separator_passed = True
                continue
            
            # If we're past the separator, parse data rows
            if in_table and separator_passed and '|' in line_stripped and line_stripped.strip():
                # Skip if this looks like another header row
                if re.search(r'Test\s+Case\s+Summary|Summary.*?Steps', line_stripped, re.IGNORECASE):
                    continue
                
                # Parse table row - split by | and extract cells
                cells = [cell.strip() for cell in line_stripped.split('|')]
                # Remove empty first/last cells if present
                if cells and not cells[0]:
                    cells = cells[1:]
                if cells and not cells[-1]:
                    cells = cells[:-1]
                
                # Should have at least 4 columns: Summary, Test Steps, Expected Result, Priority
                # But also check that first cell isn't just a separator pattern
                if len(cells) >= 4 and not re.match(r'^[\s\-:]+$', cells[0]):
                    summary = cells[0].strip()
                    test_steps = cells[1].strip() if len(cells) > 1 else ''
                    expected_result = cells[2].strip() if len(cells) > 2 else ''
                    priority = cells[3].strip() if len(cells) > 3 else 'Medium'
                    
                    # Clean up HTML tags like <br> and replace with newlines
                    test_steps = re.sub(r'<br\s*/?>', '\n', test_steps, flags=re.IGNORECASE)
                    expected_result = re.sub(r'<br\s*/?>', '\n', expected_result, flags=re.IGNORECASE)
                    
                    # Remove markdown formatting (bold, etc.)
                    summary = re.sub(r'\*\*([^*]+)\*\*', r'\1', summary)
                    test_steps = re.sub(r'\*\*([^*]+)\*\*', r'\1', test_steps)
                    expected_result = re.sub(r'\*\*([^*]+)\*\*', r'\1', expected_result)
                    
                    # Validate priority
                    if priority not in ['High', 'Medium', 'Low']:
                        priority = 'Medium'
                    
                    if summary and summary.strip():
                        # Try to extract requirement ID from summary if not found in context
                        req_id_to_use = current_req_id
                        if not req_id_to_use:
                            req_in_summary = re.search(r'REQ-?(\d+)', summary, re.IGNORECASE)
                            if req_in_summary:
                                req_id_to_use = req_in_summary.group(0)
                        
                        recommendations.append({
                            'requirement_id': req_id_to_use or 'REQ-UNKNOWN',
                            'summary': summary,
                            'test_steps': test_steps,
                            'expected_result': expected_result,
                            'priority': priority,
                            'type': 'Recommended',
                            'source': 'Gap Analysis'
                        })
            
            # Reset table state if we hit a new major section heading (but not if it's just another gap subheading)
            if line_stripped.startswith('##') and not line_stripped.startswith('###'):
                if not re.search(r'Recommendations?\s+for\s+New\s+Test\s+Cases?', line_stripped, re.IGNORECASE):
                    in_table = False
                    header_found = False
                    separator_passed = False
                    current_req_id = None
            
        
        # If we found recommendations from tables, skip other parsing
        if recommendations:
            # Clean up: Remove test_case_id from final output (not part of standard test case format)
            for rec in recommendations:
                rec.pop('test_case_id', None)
            
            # Filter out invalid recommendations (must have summary)
            recommendations = [r for r in recommendations if r.get('summary') and r.get('summary').strip()]
            
            logger.info(f"Parsed {len(recommendations)} valid recommendations from markdown tables")
            return recommendations
        
        # Pattern 1: Look for structured test cases with Test Case ID, Summary, Test Steps, Expected Result
        # Format: Test Case ID: TC-XXX-XXX followed by structured content
        # This pattern matches complete test case blocks
        tc_block_pattern = r'(?:^|\n)(?:####\s*)?(?:Gap\s+\d+[:\s]*)?.*?(?:REQ-[0-9-]+[^:]*)?\s*\n.*?Test\s+Case\s+ID[:\s]+([A-Z0-9_-]+).*?\n.*?Summary[:\s]+(.+?)(?:\n.*?Test\s+Steps[:\s]+(.+?))?(?:\n.*?Expected\s+Result[:\s]+(.+?))?(?:\n.*?Priority[:\s]+(\w+))?(?=\n(?:####|Test\s+Case\s+ID|$))'
        
        # Try to find complete test case blocks
        lines = rec_section.split('\n')
        current_tc = {}
        in_tc_block = False
        current_field = None
        current_req_id = None  # Track requirement ID from gap header
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Extract requirement ID from gap headers like "#### Gap 1: REQ-007 (Concurrent Users)"
            gap_header_match = re.search(r'Gap\s+\d+[:\s]+.*?REQ-?(\d+)', line_stripped, re.IGNORECASE)
            if gap_header_match:
                current_req_id = 'REQ-' + gap_header_match.group(1)
                continue
            
            # Skip introductory lines
            if re.search(r'^To address the identified gaps|^the following new test cases', line_stripped, re.IGNORECASE):
                continue
            
            # Look for "Test Case ID:" pattern - start of a new test case
            tc_id_match = re.search(r'Test\s+Case\s+ID[:\s]+([A-Z0-9_-]+)', line_stripped, re.IGNORECASE)
            if tc_id_match:
                # Save previous test case if it exists and has required fields
                if current_tc and current_tc.get('summary') and current_tc.get('summary').strip():
                    recommendations.append(current_tc)
                
                # Start new test case
                current_tc = {
                    'test_case_id': tc_id_match.group(1).strip(),
                    'requirement_id': current_req_id or 'REQ-UNKNOWN',
                    'summary': '',
                    'test_steps': '',
                    'expected_result': '',
                    'priority': 'Medium',
                    'type': 'Recommended',
                    'source': 'Gap Analysis'
                }
                in_tc_block = True
                current_field = None
                continue
            
            # Look for numbered test case pattern: "1. New Test Case for REQ-XXX" or "1. New Test Case"
            # This is a new format that doesn't include Test Case ID
            numbered_tc_match = re.search(r'^\d+\.\s+(?:New\s+)?Test\s+Case(?:\s+for\s+)?.*?(REQ-?\d+)?', line_stripped, re.IGNORECASE)
            if numbered_tc_match:
                # Save previous test case if it exists and has required fields
                if current_tc and current_tc.get('summary') and current_tc.get('summary').strip():
                    recommendations.append(current_tc)
                
                # Extract REQ ID from the line if present
                req_match_in_line = re.search(r'REQ-?(\d+)', line_stripped, re.IGNORECASE)
                extracted_req_id = None
                if req_match_in_line:
                    extracted_req_id = req_match_in_line.group(0)
                
                # Start new test case (without test_case_id since it's not in this format)
                current_tc = {
                    'requirement_id': extracted_req_id or current_req_id or 'REQ-UNKNOWN',
                    'summary': '',
                    'test_steps': '',
                    'expected_result': '',
                    'priority': 'Medium',
                    'type': 'Recommended',
                    'source': 'Gap Analysis'
                }
                in_tc_block = True
                current_field = None
                # Update current_req_id if we found one in this line
                if extracted_req_id:
                    current_req_id = extracted_req_id
                continue
            
            if not in_tc_block:
                continue
            
            # Extract Summary
            summary_match = re.search(r'Summary[:\s]+(.+)', line_stripped, re.IGNORECASE)
            if summary_match:
                current_tc['summary'] = summary_match.group(1).strip()
                current_field = None
                continue
            
            # Extract Test Steps (can be multi-line)
            if re.match(r'Test\s+Steps[:\s]*', line_stripped, re.IGNORECASE):
                current_field = 'test_steps'
                # Get content after "Test Steps:"
                step_content = re.sub(r'Test\s+Steps[:\s]*', '', line_stripped, flags=re.IGNORECASE).strip()
                if step_content:
                    current_tc['test_steps'] = step_content
                continue
            
            # Extract Expected Result (can be multi-line)
            if re.match(r'Expected\s+Result[:\s]*', line_stripped, re.IGNORECASE):
                current_field = 'expected_result'
                # Get content after "Expected Result:"
                exp_content = re.sub(r'Expected\s+Result[:\s]*', '', line_stripped, flags=re.IGNORECASE).strip()
                if exp_content:
                    current_tc['expected_result'] = exp_content
                continue
            
            # Extract Priority
            priority_match = re.search(r'Priority[:\s]+(\w+)', line_stripped, re.IGNORECASE)
            if priority_match:
                priority = priority_match.group(1).strip()
                if priority in ['High', 'Medium', 'Low']:
                    current_tc['priority'] = priority
                current_field = None
                continue
            
            # Extract Requirement ID from context if not already set from gap header
            if current_tc.get('requirement_id') == 'REQ-UNKNOWN' or not current_tc.get('requirement_id'):
                req_match = re.search(r'REQ-?(\d+)', line_stripped, re.IGNORECASE)
                if req_match:
                    current_tc['requirement_id'] = req_match.group(0)
            
            # Continue reading multi-line fields
            # Handle multi-line content for test_steps and expected_result
            if current_field and line_stripped:
                # Stop if we hit a new field marker or numbered test case
                if (re.match(r'Summary[:\s]', line_stripped, re.IGNORECASE) or
                    re.match(r'Test\s+Steps[:\s]', line_stripped, re.IGNORECASE) or
                    re.match(r'Expected\s+Result[:\s]', line_stripped, re.IGNORECASE) or
                    re.match(r'Priority[:\s]', line_stripped, re.IGNORECASE) or
                    re.match(r'^\d+\.\s+(?:New\s+)?Test\s+Case', line_stripped, re.IGNORECASE)):
                    # Hit a new field, stop collecting for current field
                    current_field = None
                    # Re-process this line to extract the new field
                    continue
                # Append to current field (skip markdown formatting and empty lines)
                if not line_stripped.startswith('**') and not re.match(r'^#{1,4}\s', line_stripped):
                    if current_tc.get(current_field):
                        current_tc[current_field] += '\n' + line_stripped
                    else:
                        current_tc[current_field] = line_stripped
            elif current_field and not line_stripped:
                # Empty line - keep field active but add spacing if content exists
                if current_tc.get(current_field) and not current_tc[current_field].endswith('\n'):
                    current_tc[current_field] += '\n'
        
        # Save last test case if exists
        if current_tc and current_tc.get('summary') and current_tc.get('summary').strip():
            recommendations.append(current_tc)
        
        # If no structured test cases found, try alternative pattern
        if not recommendations:
            # Alternative: Look for "To address REQ-XXX" followed by test case details
            address_pattern = r'(?:To address|For|Addressing)\s+([A-Z0-9_-]+).*?Test\s+Case\s+ID[:\s]+([A-Z0-9_-]+).*?\n.*?Summary[:\s]+(.+?)(?:\n.*?Test\s+Steps[:\s]+(.+?))?(?:\n.*?Expected\s+Result[:\s]+(.+?))?(?=\n(?:To address|For|Addressing|$))'
            
            matches = re.finditer(address_pattern, rec_section, re.DOTALL | re.IGNORECASE)
            
            for match in matches:
                req_id = match.group(1).strip()
                tc_id = match.group(2).strip()
                summary = match.group(3).strip() if match.group(3) else ''
                test_steps = match.group(4).strip() if match.group(4) else ''
                expected_result = match.group(5).strip() if match.group(5) else ''
                
                if summary and tc_id:
                    recommendations.append({
                        'test_case_id': tc_id,
                        'requirement_id': req_id,
                        'summary': summary,
                        'test_steps': test_steps,
                        'expected_result': expected_result,
                        'priority': 'Medium',
                        'type': 'Recommended',
                        'source': 'Gap Analysis'
                    })
        
        # Clean up: Remove test_case_id from final output (not part of standard test case format)
        for rec in recommendations:
            rec.pop('test_case_id', None)
        
        # Filter out invalid recommendations (must have summary)
        recommendations = [r for r in recommendations if r.get('summary') and r.get('summary').strip()]
        
        logger.info(f"Parsed {len(recommendations)} valid recommendations from gap analysis")
        
        if len(recommendations) == 0:
            logger.warning("No valid test cases found. Gap analysis text preview:")
            logger.warning(rec_section[:1000])
        
    except Exception as e:
        logger.error(f"Error parsing recommendations from gap analysis: {e}")
        logger.error(f"Exception type: {type(e).__name__}")
        import traceback
        logger.error(traceback.format_exc())
    
    return recommendations

def extract_and_save_recommended_test_cases():
    """
    Generate recommended test cases using an agent based on gap_analysis and existing test_cases.
    Reads gap_analysis and test_cases from database for current session, then uses an agent to generate new test cases.
    """
    try:
        # Check if recommended_test_cases is empty
        recommended_test_cases = st.session_state.get('recommended_test_cases', [])
        if recommended_test_cases and len(recommended_test_cases) > 0:
            logger.info("Recommended test cases already exist, skipping extraction")
            return False, "Recommended test cases already exist"
        
        # Check if ADK and runners are available
        if not ADK_AVAILABLE:
            return False, "ADK is not available. Cannot generate recommended test cases."
        
        if not runners_initialized and not initialize_runners():
            return False, "Failed to initialize agent runners."
        
        # Get session ID
        session_id = st.session_state.get(ADK_SESSION_KEY)
        if not session_id:
            return False, "Session ID not found"
        
        # Try to get gap_analysis and test_cases from session state first
        gap_analysis_text = st.session_state.get('gap_analysis', '')
        existing_test_cases = st.session_state.get('test_cases', [])
        
        # If not in session state, try to load from database
        if db:
            session_data = db.get_session_data(session_id)
            if session_data:
                if not gap_analysis_text and session_data.get('gap_analysis'):
                    gap_analysis_text = session_data['gap_analysis']
                    st.session_state.gap_analysis = gap_analysis_text
                    logger.info("Loaded gap_analysis from database")
                
                if not existing_test_cases and session_data.get('test_cases'):
                    existing_test_cases = session_data['test_cases']
                    st.session_state.test_cases = existing_test_cases
                    logger.info("Loaded test_cases from database")
        
        if not gap_analysis_text or not gap_analysis_text.strip():
            return False, "No gap analysis found. Please run gap analysis first."
        
        if not existing_test_cases or len(existing_test_cases) == 0:
            return False, "No existing test cases found. Please generate test cases first."
        
        # Prepare prompt for the agent
        agent_prompt = json.dumps({
            "gap_analysis": gap_analysis_text,
            "existing_test_cases": existing_test_cases
        })
        
        # Call the recommended test case generator agent
        logger.info("Calling recommended_test_case_generator_agent to generate new test cases")
        with SimpleLoadingOverlay(
            message="Generating Recommended Test Cases",
            submessage="Agent is analyzing gaps and creating new test cases..."
        ):
            agent_response = run_adk_sync(
                recommended_test_case_generator_runner,
                session_id,
                agent_prompt
            )
        
        # Check for error messages
        if agent_response.startswith("Error:"):
            logger.error(f"Agent returned error: {agent_response}")
            return False, f"Agent error: {agent_response}"
        
        if not agent_response or not agent_response.strip():
            return False, "Agent returned empty response"
        
        # Parse JSON response
        try:
            response_data = extract_json_from_response(agent_response)
        except (json.JSONDecodeError, ValueError) as json_error:
            logger.error(f"JSON parsing failed. Response: {agent_response[:500]}...")
            return False, f"Failed to parse agent response: {json_error}"
        
        # Validate response structure
        if 'test_cases' not in response_data:
            logger.error(f"Invalid response format. Expected 'test_cases' key. Got: {list(response_data.keys())}")
            return False, f"Invalid response format. Expected 'test_cases' key, got: {list(response_data.keys())}"
        
        recommendations = response_data['test_cases']
        
        if not recommendations or len(recommendations) == 0:
            return False, "Agent generated no test cases"
        
        # Validate each recommendation has required fields
        for idx, rec in enumerate(recommendations):
            if not rec.get('summary') or not rec.get('summary').strip():
                logger.warning(f"Recommendation {idx} missing summary, skipping")
                continue
        
        # Filter out invalid recommendations
        recommendations = [r for r in recommendations if r.get('summary') and r.get('summary').strip()]
        
        if not recommendations:
            return False, "No valid test cases generated"
        
        # Save to session state
        st.session_state.recommended_test_cases = recommendations
        
        # Initialize editing state for recommended test cases
        if 'edited_recommendations' not in st.session_state:
            st.session_state.edited_recommendations = {}
        if 'included_recommendations' not in st.session_state:
            st.session_state.included_recommendations = {}
        
        for idx in range(len(recommendations)):
            if idx not in st.session_state.edited_recommendations:
                st.session_state.edited_recommendations[idx] = recommendations[idx].copy()
            if idx not in st.session_state.included_recommendations:
                st.session_state.included_recommendations[idx] = True
        
        # Save to database
        if db:
            complete_data = {
                'requirements': st.session_state.context.get('requirements', []),
                'context': st.session_state.context.get('context', ''),
                'requirement_contexts': st.session_state.context.get('requirement_contexts', {}),
                'test_cases': existing_test_cases,
                'gap_analysis': gap_analysis_text,
                'jira_csv': st.session_state.get('jira_csv', ''),
                'recommended_test_cases': recommendations,
                'included_recommendations': st.session_state.get('included_recommendations', {}),
                'edited_recommendations': st.session_state.get('edited_recommendations', {}),
                'edited_test_cases': st.session_state.get('edited_test_cases', {}),
                'included_test_cases': st.session_state.get('included_test_cases', {}),
                'finalized_test_cases': st.session_state.get('finalized_test_cases', []),
                'is_finalized': st.session_state.get('is_finalized', False)
            }
            db.store_session_data(session_id, complete_data)
            logger.info(f"Generated and saved {len(recommendations)} recommended test cases to database")
        
        return True, f"Successfully generated {len(recommendations)} recommended test cases"
        
    except Exception as e:
        error_msg = f"Failed to extract recommended test cases: {e}"
        logger.error(error_msg)
        import traceback
        logger.error(traceback.format_exc())
        return False, error_msg

def extract_json_from_response(response: str) -> dict:
    """Extract JSON from response with better error handling."""
    if not response or not response.strip():
        raise ValueError("Empty response from agent")
    
    # Clean the response
    response = response.strip()
    
    try:
        # Try parsing entire response
        return json.loads(response)
    except json.JSONDecodeError:
        # Look for JSON within markdown code blocks
        import re
        
        # Try to find JSON in code blocks
        json_blocks = re.findall(r'```(?:json)?\s*(\{.*?\})\s*```', response, re.DOTALL)
        for block in json_blocks:
            try:
                return json.loads(block)
            except json.JSONDecodeError:
                continue
        
        # Try to find JSON by braces
        json_start = response.find('{')
        json_end = response.rfind('}') + 1
        
        if json_start != -1 and json_end > json_start:
            try:
                return json.loads(response[json_start:json_end])
            except json.JSONDecodeError:
                pass
        
        # Last resort: try to find any JSON-like structure
        for i in range(json_start, len(response)):
            for j in range(len(response), json_start, -1):
                try:
                    potential_json = response[i:j]
                    if potential_json.strip().startswith('{') and potential_json.strip().endswith('}'):
                        return json.loads(potential_json)
                except json.JSONDecodeError:
                    continue
        
        raise ValueError(f"No valid JSON found in response. Response preview: {response[:200]}...")

# --- Main Processing Functions ---
def handle_context_retrieval(uploaded_file):
    """Handle context retrieval with comprehensive error handling."""
    if not ADK_AVAILABLE:
        st.error("ADK is not available. Cannot process document.")
        return
    
    # with st.spinner("Context Retriever Agent is analyzing your document..."):
    # with LoadingAnimation(
    #     title="Analyzing Document",
    #     subtitle="AI Agent is extracting requirements...",
    #     description="This may take a few moments while we process and analyze your document for requiremenst and underline context",
    #     animation_type="dots"
    # ):
    with SimpleLoadingOverlay(
        message="Analyzing Document", 
        submessage="Agent is extracting requirements and building context..."
    ):
        try:
            # Parse file
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            parsers = {
                ".pdf": parse_pdf,
                ".docx": parse_docx,
                ".txt": parse_txt
            }
            
            parser = parsers.get(file_extension, parse_txt)
            file_content = parser(uploaded_file)
            
            if not file_content:
                st.error("Could not extract meaningful content from the document.")
                return
            
            # Prepare prompt
            prompt = f"""Here is the requirement document content:

---
{file_content}
---

Please analyze this document and extract requirements in the specified JSON format."""

            # Run agent
            #response = run_agent_with_retry(context_retriever_runner, prompt)
            print("context_retriever:")
            print("current_session_id:", st.session_state[ADK_SESSION_KEY])
            response = run_adk_sync(context_retriever_runner, st.session_state[ADK_SESSION_KEY], prompt)
            
            # Check for error messages
            if response.startswith("Error:"):
                st.error(response)
                st.session_state.last_error = response
                return
            
            if not response or not response.strip():
                st.error("Agent returned empty response.")
                st.session_state.last_error = "Agent returned empty response."
                return
            
            # Parse JSON
            try:
                context_data = extract_json_from_response(response)
            except (json.JSONDecodeError, ValueError) as json_error:
                logger.error(f"JSON parsing failed. Response: {response[:500]}...")
                st.error(f"Failed to parse agent response: {json_error}")
                
                # Show raw response for debugging
                with st.expander("Show raw agent response for debugging"):
                    st.text(response)
                return
            
            # Validate structure
            required_keys = ['requirements', 'context']
            if not all(key in context_data for key in required_keys):
                st.error(f"Invalid response structure. Expected keys: {required_keys}, got: {list(context_data.keys())}")
                return
            
            st.session_state.context = context_data
            
            # Store in database
            if db:
                try:
                    db.store_session_data(st.session_state[ADK_SESSION_KEY], context_data)
                    logger.info("Context data stored in database")
                except Exception as e:
                    logger.warning(f"Failed to store in database: {e}")
            
            st.success("Document analysis completed successfully!")
            
        except Exception as e:
            error_msg = f"Context retrieval failed: {e}"
            logger.error(error_msg)
            st.error(error_msg)
            st.session_state.last_error = error_msg

def run_generation_pipeline(edited_context, updated_context_data=None):
    """Run the complete test case generation pipeline."""
    try:
        # Use provided context_data if available, otherwise use session state
        if updated_context_data:
            st.session_state.context = updated_context_data
        else:
            # Update context in session state
            st.session_state.context['context'] = edited_context
        
        # Step 1: Generate test cases
        # with st.spinner("Generating test cases..."):
        with SimpleLoadingOverlay(
            message="Generating Test Cases", 
            submessage="Agent is analysing context and generating comprehensive test scenarios..."
        ):
            test_case_prompt = json.dumps(st.session_state.context)
            
            test_response = run_adk_sync(test_case_generator_runner, st.session_state[ADK_SESSION_KEY], test_case_prompt)
            
            # Check for error messages
            if test_response.startswith("Error:"):
                st.error(test_response)
                st.session_state.last_error = test_response
                return
            
            if not test_response or not test_response.strip():
                st.error("Test case generator returned empty response.")
                st.session_state.last_error = "Test case generator returned empty response."
                return
            
            try:
                test_data = extract_json_from_response(test_response)
            except (json.JSONDecodeError, ValueError) as json_error:
                logger.error(f"Test case JSON parsing failed. Response: {test_response[:500]}...")
                st.error(f"Failed to parse test case response: {json_error}")
                with st.expander("Show raw test case generator response"):
                    st.text(test_response)
                st.session_state.last_error = f"Test case JSON parsing failed: {json_error}"
                return
            
            if 'test_cases' not in test_data:
                st.error(f"Invalid test case response format. Expected 'test_cases' key. Got: {list(test_data.keys())}")
                with st.expander("Show raw test case response"):
                    st.json(test_data)
                st.session_state.last_error = "Invalid test case response format"
                return
            
            st.session_state.test_cases = test_data['test_cases']
            print("Test cases generated successfully!")
        
        # Step 2: Compliance review
        # with st.spinner("Reviewing for compliance..."):
        with SimpleLoadingOverlay(
            message="Compliance Review", 
            submessage="Agent is validating generated test cases against Healthcare regulations..."
        ):
            compliance_prompt = json.dumps({
                "test_cases": st.session_state.test_cases,
                "context": st.session_state.context['context']
            })
            
            compliance_response = run_adk_sync(compliance_runner, st.session_state[ADK_SESSION_KEY], compliance_prompt)
            
            # Check for error messages
            if compliance_response.startswith("Error:"):
                logger.error(f"Compliance agent returned error: {compliance_response}")
                st.error(compliance_response)
                st.session_state.last_error = compliance_response
                return
            
            # Check if response contains error indicators
            if "UNEXPECTED_TOOL_CALL" in compliance_response or ("FinishReason" in compliance_response and "UNEXPECTED" in compliance_response):
                error_msg = "Compliance agent encountered a tool call error (RAG tool issue). Retrying without RAG..."
                logger.warning(f"{error_msg} Response: {compliance_response[:200]}")
                
                # Try to continue without RAG - update the agent to not use tools temporarily
                # For now, show error but allow user to proceed
                st.warning("⚠️ RAG tool error detected. The compliance agent tried to use the RAG tool but it's not properly configured.")
                st.info("💡 **Workaround**: The compliance agent can work with the provided context without RAG. You can:")
                st.markdown("""
                1. **Continue without RAG**: The compliance review may still work with the context you provided
                2. **Fix RAG tool**: Check that chromadb and sentence-transformers are properly installed
                3. **Check logs**: See detailed error information below
                """)
                
                with st.expander("🔧 Technical Details", expanded=False):
                    st.text(compliance_response[:1000])
                    st.caption("This error means the RAG tool couldn't be called. The agent can still work with provided context.")
                
                # Don't block - let the user decide
                st.error("❌ Cannot proceed: Compliance agent needs to be fixed. The RAG tool configuration issue prevents completion.")
                st.session_state.last_error = error_msg
                return
            
            if not compliance_response or not compliance_response.strip():
                logger.error("Compliance agent returned empty response")
                st.error("Compliance agent returned empty response.")
                st.session_state.last_error = "Compliance agent returned empty response."
                return
            
            try:
                compliance_data = extract_json_from_response(compliance_response)
                logger.info(f"Compliance response parsed successfully. Keys: {list(compliance_data.keys())}")
            except (json.JSONDecodeError, ValueError) as json_error:
                logger.error(f"Compliance JSON parsing failed. Response length: {len(compliance_response)} chars")
                logger.error(f"Compliance JSON parsing failed. Response preview: {compliance_response[:500]}...")
                logger.error(f"JSON Error: {json_error}")
                st.error(f"Failed to parse compliance agent response: {json_error}")
                with st.expander("Show raw compliance agent response"):
                    st.text(compliance_response)
                    st.caption(f"Response length: {len(compliance_response)} characters")
                st.session_state.last_error = f"Compliance JSON parsing failed: {json_error}"
                return
            
            # Log what we actually got
            logger.info(f"Compliance response structure - Keys: {list(compliance_data.keys())}")
            logger.info(f"Compliance response type: {type(compliance_data)}")
            
            if 'compliant_test_cases' not in compliance_data:
                # Detailed logging of what we got
                logger.error(f"⚠️ Compliance response missing 'compliant_test_cases' key!")
                logger.error(f"Response keys: {list(compliance_data.keys())}")
                logger.error(f"Full compliance response (first 1500 chars): {compliance_response[:1500]}")
                logger.error(f"Parsed compliance_data type: {type(compliance_data)}")
                logger.error(f"Parsed compliance_data (first 1000 chars): {str(compliance_data)[:1000]}")
                
                # Check if test cases might be directly in the response or under a different key
                if isinstance(compliance_data, dict):
                    for key in compliance_data.keys():
                        logger.error(f"  Key '{key}' type: {type(compliance_data[key])}")
                        if isinstance(compliance_data[key], list):
                            logger.error(f"  Key '{key}' is a list with {len(compliance_data[key])} items")
                        elif isinstance(compliance_data[key], dict):
                            logger.error(f"  Key '{key}' is a dict with keys: {list(compliance_data[key].keys())[:10]}")
                
                st.error(f"Invalid compliance response format. Expected 'compliant_test_cases' key. Got: {list(compliance_data.keys())}")
                with st.expander("Show raw compliance response (debug)", expanded=True):
                    st.subheader("Raw Response Text")
                    st.text_area("Full response:", value=compliance_response, height=400, key="compliance_raw_response")
                    st.subheader("Parsed JSON Structure")
                    st.json(compliance_data)
                    st.subheader("Response Analysis")
                    st.write(f"**Response Length:** {len(compliance_response)} characters")
                    st.write(f"**Parsed Keys:** {list(compliance_data.keys())}")
                    if isinstance(compliance_data, dict):
                        for key, value in compliance_data.items():
                            st.write(f"- **{key}**: {type(value).__name__}" + 
                                   (f" (length: {len(value)})" if isinstance(value, (list, str, dict)) else ""))
                st.session_state.last_error = "Invalid compliance response format"
                return
            
            st.session_state.test_cases = compliance_data['compliant_test_cases']
            print("Compliance review completed successfully!")
        
        # Step 3: Gap analysis
        # with st.spinner("Performing gap analysis..."):
        with SimpleLoadingOverlay(
            message="Feature Gap Analysis", 
            submessage="Agent is performing feature gap analysis for generated test cases..."
        ):
            gap_prompt = json.dumps({
                "requirements": st.session_state.context['requirements'],
                "test_cases": st.session_state.test_cases
            })

            gap_response = run_adk_sync(gap_analysis_runner, st.session_state[ADK_SESSION_KEY], gap_prompt)
            st.session_state.gap_analysis = gap_response
            
            # Parse recommendations from gap analysis
            st.session_state.recommended_test_cases = parse_recommendations_from_gap_analysis(gap_response)
            logger.info(f"Parsed {len(st.session_state.recommended_test_cases)} recommended test cases from gap analysis")
            
            print("Gap analysis completed successfully!")
        
        # Step 4: Generate Jira CSV
        # with st.spinner("Formatting for Jira..."):
        with SimpleLoadingOverlay(
            message="Generate JIRA formatted test cases", 
            submessage="Agent is transforming test cases in to JIRA format for easy import..."
        ):
            jira_prompt = json.dumps({"test_cases": st.session_state.test_cases})
            jira_response = run_adk_sync(jira_formatter_runner, st.session_state[ADK_SESSION_KEY], jira_prompt)
            st.session_state.jira_csv = jira_response
        
        st.session_state.processing_complete = True
        
        # Save all data to database after pipeline completes
        if db:
            try:
                # Prepare complete session data
                complete_data = {
                    'requirements': st.session_state.context.get('requirements', []),
                    'context': st.session_state.context.get('context', ''),
                    'requirement_contexts': st.session_state.context.get('requirement_contexts', {}),
                    'test_cases': st.session_state.test_cases,
                    'gap_analysis': st.session_state.gap_analysis,
                    'jira_csv': st.session_state.jira_csv,
                    'recommended_test_cases': st.session_state.recommended_test_cases,
                    'included_recommendations': st.session_state.get('included_recommendations', {}),
                    'edited_recommendations': st.session_state.get('edited_recommendations', {}),
                    'edited_test_cases': st.session_state.get('edited_test_cases', {}),
                    'included_test_cases': st.session_state.get('included_test_cases', {}),
                    'finalized_test_cases': st.session_state.get('finalized_test_cases', []),
                    'is_finalized': st.session_state.get('is_finalized', False)
                }
                
                db.store_session_data(st.session_state[ADK_SESSION_KEY], complete_data)
                logger.info("Complete session data saved to database after pipeline completion")
            except Exception as e:
                logger.warning(f"Failed to save complete session data to database: {e}")
        
        st.success("All processing completed successfully!")
        
    except Exception as e:
        error_msg = f"Pipeline execution failed: {e}"
        logger.error(error_msg)
        st.error(error_msg)
        st.session_state.last_error = error_msg

    
def generate_test_case_id(index: int) -> str:
    """Generate a test case ID in the format TC-001, TC-002, etc."""
    return f"TC-{index + 1:03d}"

def finalize_test_cases():
    """Collect all included and edited test cases into finalized list and save."""
    finalized = []
    
    # Add included generated test cases (with edits)
    if st.session_state.test_cases:
        for idx, tc in enumerate(st.session_state.test_cases):
            if st.session_state.included_test_cases.get(idx, True):
                # Get edited version if exists, otherwise use original
                edited = st.session_state.edited_test_cases.get(idx, tc.copy())
                finalized.append(edited.copy())
    
    # Add included recommended test cases (with edits)
    if st.session_state.get('recommended_test_cases'):
        for idx in range(len(st.session_state.recommended_test_cases)):
            if st.session_state.included_recommendations.get(idx, True):
                edited_rec = st.session_state.edited_recommendations.get(idx, st.session_state.recommended_test_cases[idx].copy())
                # Create clean test case from edited recommendation
                new_test_case = {
                    'requirement_id': edited_rec.get('requirement_id', 'REQ-UNKNOWN'),
                    'summary': edited_rec.get('summary', ''),
                    'test_steps': edited_rec.get('test_steps', ''),
                    'expected_result': edited_rec.get('expected_result', ''),
                    'priority': edited_rec.get('priority', 'Medium')
                }
                # Add compliance notes if present
                if edited_rec.get('compliance_notes'):
                    new_test_case['compliance_notes'] = edited_rec.get('compliance_notes')
                
                # Remove empty fields
                new_test_case = {k: v for k, v in new_test_case.items() if v}
                
                # Add to finalized if summary exists
                if new_test_case.get('summary'):
                    finalized.append(new_test_case)
    
    # Remove duplicates based on summary
    seen_summaries = set()
    unique_finalized = []
    for tc in finalized:
        summary = tc.get('summary', '')
        if summary and summary not in seen_summaries:
            unique_finalized.append(tc)
            seen_summaries.add(summary)
    
    # Add test case IDs to finalized test cases
    for idx, tc in enumerate(unique_finalized):
        if 'test_case_id' not in tc or not tc.get('test_case_id'):
            tc['test_case_id'] = generate_test_case_id(idx)
    
    # Save to session state
    st.session_state.finalized_test_cases = unique_finalized
    st.session_state.is_finalized = True
    
    # Save to database
    if db:
        try:
            complete_data = {
                'requirements': st.session_state.context.get('requirements', []),
                'context': st.session_state.context.get('context', ''),
                'requirement_contexts': st.session_state.context.get('requirement_contexts', {}),
                'test_cases': st.session_state.test_cases,
                'gap_analysis': st.session_state.gap_analysis,
                'jira_csv': st.session_state.jira_csv,
                'recommended_test_cases': st.session_state.recommended_test_cases,
                'included_recommendations': st.session_state.get('included_recommendations', {}),
                'edited_recommendations': st.session_state.get('edited_recommendations', {}),
                'edited_test_cases': st.session_state.get('edited_test_cases', {}),
                'included_test_cases': st.session_state.get('included_test_cases', {}),
                'finalized_test_cases': unique_finalized,
                'is_finalized': True
            }
            db.store_session_data(st.session_state[ADK_SESSION_KEY], complete_data)
            logger.info("Finalized test cases saved to database")
            return True
        except Exception as e:
            logger.error(f"Failed to save finalized test cases to database: {e}")
            return False
    
    return True


def display_results():
    """Display the generated results with unified test cases view."""
    # Check if test cases are finalized
    if st.session_state.get('is_finalized') and st.session_state.get('finalized_test_cases'):
        # Show finalized test cases view (read-only)
        display_finalized_test_cases()
        return
    
    # Show unified test cases list with editing capabilities
    st.header("📋 All Test Cases")
    
    # Initialize state variables
    if 'edited_test_cases' not in st.session_state:
        st.session_state.edited_test_cases = {}
    if 'included_test_cases' not in st.session_state:
        st.session_state.included_test_cases = {}
        if 'included_recommendations' not in st.session_state:
            st.session_state.included_recommendations = {}
        if 'edited_recommendations' not in st.session_state:
            st.session_state.edited_recommendations = {}
        
    # Initialize generated test cases editing state
    if st.session_state.test_cases:
        for idx in range(len(st.session_state.test_cases)):
            if idx not in st.session_state.edited_test_cases:
                st.session_state.edited_test_cases[idx] = st.session_state.test_cases[idx].copy()
            # Default to included (True) if not set
            if idx not in st.session_state.included_test_cases:
                st.session_state.included_test_cases[idx] = True
    
    # Initialize recommended test cases editing state
    if st.session_state.get('recommended_test_cases'):
        for idx in range(len(st.session_state.recommended_test_cases)):
            if idx not in st.session_state.edited_recommendations:
                st.session_state.edited_recommendations[idx] = st.session_state.recommended_test_cases[idx].copy()
            # Default to included (True) if not set
            if idx not in st.session_state.included_recommendations:
                st.session_state.included_recommendations[idx] = True
        
    # Calculate counts - be explicit about checking recommended_test_cases
    total_generated = len(st.session_state.test_cases) if st.session_state.test_cases else 0
    recommended_test_cases_list = st.session_state.get('recommended_test_cases', [])
    # Ensure it's a list and not None or empty string
    if not isinstance(recommended_test_cases_list, list):
        recommended_test_cases_list = []
    total_recommended = len(recommended_test_cases_list)
    included_generated = sum(1 for v in st.session_state.included_test_cases.values() if v)
    included_recommended = sum(1 for v in st.session_state.included_recommendations.values() if v)
    
    # Header with summary and finalize button
    col1, col2 = st.columns([3, 1])
    with col1:
        st.info(f"📊 **Total Test Cases:** {total_generated + total_recommended} (Generated: {total_generated}, Recommended: {total_recommended}) | **Included:** {included_generated + included_recommended}")
    with col2:
        if st.button("💾 Finalize Test Cases", type="primary", use_container_width=True):
            if finalize_test_cases():
                st.success("✅ Test cases finalized successfully!")
                st.rerun()
            else:
                st.error("❌ Failed to finalize test cases. Please try again.")
    
    st.divider()
    
    # Display all test cases in unified list
    all_test_cases = []
    
    # Add generated test cases with badge
    if st.session_state.test_cases:
        for idx, tc in enumerate(st.session_state.test_cases):
            all_test_cases.append({
                'index': idx,
                'type': 'generated',
                'data': tc
            })
    
    # Add recommended test cases with badge - check explicitly for list with items
    if recommended_test_cases_list and len(recommended_test_cases_list) > 0:
        for idx, rec in enumerate(recommended_test_cases_list):
            # Ensure rec is a dictionary
            if isinstance(rec, dict):
                all_test_cases.append({
                    'index': idx,
                    'type': 'recommended',
                    'data': rec
                })
    elif st.session_state.get('gap_analysis') and total_recommended == 0:
        # Show warning if gap analysis exists but no recommendations were parsed
        st.warning("⚠️ Gap analysis found but no recommended test cases were parsed. Check gap analysis section below.")
    
    # Display unified list
    if all_test_cases:
        for item in all_test_cases:
            idx = item['index']
            tc_type = item['type']
            tc_data = item['data']
            
            # Get edited version
            if tc_type == 'generated':
                edited_tc = st.session_state.edited_test_cases.get(idx, tc_data.copy())
                is_included = st.session_state.included_test_cases.get(idx, True)
                badge = "🟢 Generated"
                toggle_key = f"gen_include_{idx}"
            else:  # recommended
                edited_tc = st.session_state.edited_recommendations.get(idx, tc_data.copy())
                is_included = st.session_state.included_recommendations.get(idx, True)
                badge = "🔵 Recommended"
                toggle_key = f"rec_include_{idx}"
            
            # Create columns: Badge & Toggle | Editable Content
            col_badge_toggle, col_content = st.columns([0.2, 0.8])
            
            with col_badge_toggle:
                st.markdown(f"**{badge}**")
                
                # Include/Exclude toggle
                toggle_value = st.toggle(
                    "Include" if is_included else "Exclude",
                    value=is_included,
                    key=toggle_key,
                    label_visibility="visible"
                )
                
                # Update state
                if tc_type == 'generated':
                    st.session_state.included_test_cases[idx] = toggle_value
                else:
                    st.session_state.included_recommendations[idx] = toggle_value
                
                if toggle_value:
                    st.success("✓ Included")
                else:
                    st.warning("✗ Excluded")
            
            with col_content:
                summary_preview = edited_tc.get('summary', 'N/A')[:70]
                req_id = edited_tc.get('requirement_id', 'N/A')
                with st.expander(
                    f"**{summary_preview}...** - {req_id}",
                    expanded=True
                ):
                    # Editable fields
                    col_a, col_b = st.columns(2)
                    with col_a:
                        edited_tc['requirement_id'] = st.text_input(
                            "Requirement ID",
                            value=edited_tc.get('requirement_id', ''),
                            key=f"{tc_type}_req_id_{idx}"
                        )
                    
                    with col_b:
                        priority_options = ['High', 'Medium', 'Low']
                        current_priority = edited_tc.get('priority', 'Medium')
                        priority_index = priority_options.index(current_priority) if current_priority in priority_options else 1
                        edited_tc['priority'] = st.selectbox(
                            "Priority",
                            options=priority_options,
                            index=priority_index,
                            key=f"{tc_type}_priority_{idx}"
                        )
                    
                    edited_tc['summary'] = st.text_input(
                        "Test Case Summary",
                        value=edited_tc.get('summary', ''),
                        key=f"{tc_type}_summary_{idx}"
                    )
                    
                    if edited_tc.get('description'):
                        st.write(f"**Original Description (for reference):**")
                        st.write(edited_tc.get('description', ''))
                    
                    edited_tc['test_steps'] = st.text_area(
                        "Test Steps",
                        value=edited_tc.get('test_steps', ''),
                        height=150,
                        key=f"{tc_type}_steps_{idx}",
                        help="Enter the detailed test steps for this test case"
                    )
                    
                    edited_tc['expected_result'] = st.text_area(
                        "Expected Result",
                        value=edited_tc.get('expected_result', ''),
                        height=100,
                        key=f"{tc_type}_expected_{idx}",
                        help="Enter the expected result for this test case"
                    )
                    
                    # Show compliance notes if present (read-only for generated, editable for recommended)
                    if edited_tc.get('compliance_notes'):
                        if tc_type == 'generated':
                            st.write(f"**Compliance Notes:**")
                            st.write(edited_tc.get('compliance_notes'))
                        else:
                            edited_tc['compliance_notes'] = st.text_area(
                                "Compliance Notes",
                                value=edited_tc.get('compliance_notes', ''),
                                height=100,
                                key=f"{tc_type}_compliance_{idx}"
                            )
                    
                    # Update edited state
                    if tc_type == 'generated':
                        st.session_state.edited_test_cases[idx] = edited_tc
                    else:
                        st.session_state.edited_recommendations[idx] = edited_tc
        
        st.divider()
    
    # Gap Analysis
    if st.session_state.gap_analysis:
        st.subheader("📉 Gap Analysis Report")
        st.markdown(st.session_state.gap_analysis)
        
        # Add extraction button if recommended_test_cases is empty
        recommended_test_cases = st.session_state.get('recommended_test_cases', [])
        if not recommended_test_cases or len(recommended_test_cases) == 0:
            col1, col2 = st.columns([3, 1])
            with col1:
                st.info("💡 Recommended test cases not found. Click the button to extract them from the gap analysis.")
            with col2:
                if st.button("🔍 Extract Recommended Test Cases", type="primary", use_container_width=True, key="extract_recommendations"):
                    success, message = extract_and_save_recommended_test_cases()
                    if success:
                        st.success(message)
                        st.rerun()  # Reload the page after extraction
                    else:
                        st.error(message)
        else:
            st.info(f"✅ {len(recommended_test_cases)} recommended test cases already extracted")
        
        st.divider()
    
    # Actions sidebar
    with st.sidebar:
        st.header("🎬 Actions")
        
        # Download current test cases (non-finalized, includes edited versions)
        all_current_test_cases = []
        
        # Add included generated test cases (with edits)
        if st.session_state.test_cases:
            for idx in range(len(st.session_state.test_cases)):
                if st.session_state.included_test_cases.get(idx, True):
                    edited = st.session_state.edited_test_cases.get(idx, st.session_state.test_cases[idx].copy())
                    all_current_test_cases.append(edited.copy())
        
        # Add included recommended test cases (with edits)
        if st.session_state.get('recommended_test_cases'):
            for idx in range(len(st.session_state.recommended_test_cases)):
                if st.session_state.included_recommendations.get(idx, True):
                    edited_rec = st.session_state.edited_recommendations.get(idx, st.session_state.recommended_test_cases[idx].copy())
                    new_test_case = {
                        'requirement_id': edited_rec.get('requirement_id', 'REQ-UNKNOWN'),
                        'summary': edited_rec.get('summary', ''),
                        'test_steps': edited_rec.get('test_steps', ''),
                        'expected_result': edited_rec.get('expected_result', ''),
                        'priority': edited_rec.get('priority', 'Medium')
                    }
                    if edited_rec.get('compliance_notes'):
                        new_test_case['compliance_notes'] = edited_rec.get('compliance_notes')
                    new_test_case = {k: v for k, v in new_test_case.items() if v}
                    if new_test_case.get('summary'):
                        all_current_test_cases.append(new_test_case)
        
        # Remove duplicates
        seen_summaries = set()
        unique_test_cases = []
        for tc in all_current_test_cases:
            summary = tc.get('summary', '')
            if summary and summary not in seen_summaries:
                unique_test_cases.append(tc)
                seen_summaries.add(summary)
        
        if unique_test_cases:
            # Add test case IDs if not present
            test_cases_with_ids = []
            for idx, tc in enumerate(unique_test_cases):
                tc_copy = tc.copy()
                if 'test_case_id' not in tc_copy or not tc_copy.get('test_case_id'):
                    tc_copy['test_case_id'] = generate_test_case_id(idx)
                test_cases_with_ids.append(tc_copy)
            
            test_cases_json = json.dumps(test_cases_with_ids, indent=2)
            st.download_button(
                label=f"📥 Download Test Cases (JSON) - {len(test_cases_with_ids)} included",
                data=test_cases_json,
                file_name=f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
        
        # Start new analysis
        if st.button("🔄 Start New Analysis", use_container_width=True):
            reset_run_state()
            st.rerun()


def display_finalized_test_cases():
    """Display finalized test cases in read-only view."""
    st.header("✅ Finalized Test Cases")
    st.success(f"Test cases have been finalized! Total: {len(st.session_state.finalized_test_cases)} test cases")
    st.info("💡 These test cases are now locked and ready for download. To make changes, you'll need to start a new analysis.")
    
    # Display finalized test cases
    if st.session_state.finalized_test_cases:
        for i, tc in enumerate(st.session_state.finalized_test_cases, 1):
            with st.expander(f"Test Case {i}: {tc.get('summary', 'N/A')}", expanded=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Requirement ID:** {tc.get('requirement_id', 'N/A')}")
                with col2:
                    st.write(f"**Priority:** {tc.get('priority', 'N/A')}")
                
                st.write(f"**Test Steps:**")
                st.write(tc.get('test_steps', 'N/A'))
                
                st.write(f"**Expected Result:**")
                st.write(tc.get('expected_result', 'N/A'))
                
                if tc.get('compliance_notes'):
                    st.write(f"**Compliance Notes:**")
                    st.write(tc.get('compliance_notes'))
    else:
        st.warning("No finalized test cases available.")
    
    st.divider()
    
    # Download finalized test cases
    with st.sidebar:
        st.header("📥 Download Finalized Test Cases")
        
        if st.session_state.finalized_test_cases:
            # Ensure all finalized test cases have IDs (they should from finalize_test_cases, but double-check)
            finalized_with_ids = []
            for idx, tc in enumerate(st.session_state.finalized_test_cases):
                tc_copy = tc.copy()
                if 'test_case_id' not in tc_copy or not tc_copy.get('test_case_id'):
                    tc_copy['test_case_id'] = generate_test_case_id(idx)
                finalized_with_ids.append(tc_copy)
            
            # JSON download
            finalized_json = json.dumps(finalized_with_ids, indent=2)
            st.download_button(
                label=f"📥 Download JSON ({len(finalized_with_ids)} test cases)",
                data=finalized_json,
                file_name=f"finalized_test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
            
            # Generate and download Jira CSV
            try:
                jira_prompt = json.dumps({"test_cases": finalized_with_ids})
                finalized_jira_csv = run_adk_sync(jira_formatter_runner, st.session_state[ADK_SESSION_KEY], jira_prompt)
                
                st.download_button(
                    label=f"📥 Download Jira CSV ({len(st.session_state.finalized_test_cases)} test cases)",
                    data=finalized_jira_csv,
                    file_name=f"finalized_jira_test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            except Exception as e:
                logger.warning(f"Failed to generate Jira CSV for finalized test cases: {e}")
                st.warning("Could not generate Jira CSV for finalized test cases.")
        
        # Start new analysis
        if st.button("🔄 Start New Analysis", use_container_width=True):
            reset_run_state()
            st.rerun()


def set_sidebar_background_with_overlay(image_file):
    """Set sidebar background with semi-transparent overlay"""
    try:
        image_path = os.path.join('images', image_file)
        with open(image_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode()
        
        st.markdown(f"""
        <style>
        [data-testid="stSidebar"] {{
            background-image: linear-gradient(rgba(255,255,255,0.3), rgba(255,255,255,0.3)),
                              url('data:image/png;base64,{img_data}');
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
        }}
        </style>
        """, unsafe_allow_html=True)
    except FileNotFoundError:
        st.warning(f"Sidebar background image {image_file} not found")

def main():
    """Main application with better error handling and user feedback."""
    # Configure Streamlit page
    st.set_page_config(
        page_title="AI Test Case Generator", 
        layout="wide", 
        initial_sidebar_state="expanded"
    )


    st.title("🤖 AI-Powered Healthcare Test Case Generator")
    st.markdown("""
    Upload your software requirement specifications, and let our team of AI agents generate compliant,
    traceable, and Jira-ready test cases for you.
    """)

    #set_sidebar_background_with_overlay('img1.png')

    # Initialize session state
    initialize_session_state()
    
    # Initialize RAG service
    try:
        rag_service = get_rag_service()
        if rag_service:
            logger.info("RAG service initialized successfully")
        else:
            logger.warning("RAG service not available")
    except Exception as e:
        logger.warning(f"RAG service initialization failed: {e}")

    # System status check in sidebar
    with st.sidebar:
        st.header("⚙️ System Status")
        #st.sidebar.image('images/img1.png')
        # Use an absolute path for reliability
        logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'images', 'img1.png')
        if os.path.exists(logo_path):
            st.sidebar.image(logo_path)
        
        if ADK_AVAILABLE:
            st.success("✅ Google ADK Available")
            if initialize_runners():
                st.success("✅ Agent Runners Ready")
            else:
                st.error("❌ Agent Runners Failed")
        else:
            st.error("❌ Google ADK Not Available")
        
        if db:
            st.success("✅ Database Connected")
        else:
            st.warning("⚠️ Database Not Available")

    # Error display with log viewer
    if st.session_state.get('last_error'):
        st.error(f"⚠️ **Last Error:** {st.session_state.last_error}")
        
        # Add expander to view recent logs
        with st.expander("🔍 View Recent Logs (for debugging)", expanded=False):
            # Get current session ID
            current_session_id = st.session_state.get(ADK_SESSION_KEY, None)
            
            # Show both main log and session-specific log
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📋 Main Log File")
                log_file_path = _main_log_file_path
                
                # Also try relative path as fallback
                if not os.path.exists(log_file_path):
                    log_file_path = "healthcare_agent.log"
                
                if os.path.exists(log_file_path):
                    try:
                        # Read last 50 lines of log file
                        with open(log_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            lines = f.readlines()
                            if lines:
                                recent_lines = lines[-50:] if len(lines) > 50 else lines
                                log_content = ''.join(recent_lines)
                                st.text_area(
                                    "Last 50 log lines:",
                                    value=log_content,
                                    height=300,
                                    help=f"These are the last 50 lines from {log_file_path}",
                                    key="log_viewer"
                                )
                                st.caption(f"📁 Log file: `{log_file_path}` | Total lines: {len(lines)}")
                            else:
                                st.warning("Log file exists but is empty. Logs may be going to console only.")
                                st.info("💡 Check the Streamlit console/terminal for error messages.")
                    except Exception as log_error:
                        st.error(f"Could not read log file: {log_error}")
                        st.info(f"💡 Try checking the console output or the file manually: `{log_file_path}`")
                else:
                    st.warning(f"⚠️ Log file not found at: `{log_file_path}`")
                    st.info("💡 Logs may be going to console only. Check the terminal/console where Streamlit is running.")
            
            with col2:
                st.subheader("🎯 Session-Specific Log")
                if current_session_id:
                    session_log_file = get_session_log_file(current_session_id)
                    if os.path.exists(session_log_file):
                        try:
                            with open(session_log_file, 'r', encoding='utf-8', errors='ignore') as f:
                                lines = f.readlines()
                                if lines:
                                    recent_lines = lines[-50:] if len(lines) > 50 else lines
                                    log_content = ''.join(recent_lines)
                                    st.text_area(
                                        f"Session: {current_session_id[:20]}...",
                                        value=log_content,
                                        height=300,
                                        help=f"Session-specific log: {session_log_file}",
                                        key="session_log_viewer"
                                    )
                                    st.caption(f"📁 Session log: `{session_log_file}` | Total lines: {len(lines)}")
                                else:
                                    st.info("Session log file is empty.")
                        except Exception as log_error:
                            st.error(f"Could not read session log: {log_error}")
                    else:
                        st.info(f"💡 Session log not created yet: `{session_log_file}`")
                        st.caption(f"Session ID: {current_session_id}")
                else:
                    st.info("No active session. Session logs are created when you start an analysis.")
        
        col1, col2 = st.columns([1, 4])
        with col1:
            if st.button("Clear Error", type="secondary"):
                st.session_state.last_error = None
                st.rerun()
        with col2:
            st.info("💡 **Tip:** Check the logs above or the console for detailed error information.")

    # Check authentication configuration
    auth_config = configure_adk_authentication()
    if not auth_config:
        st.error("⚠️ Authentication not configured!")
        st.markdown("""
        **Please configure authentication in your `.env` file:**
        
        **Option 1: Use API Key (Recommended for local development)**
        ```env
        GOOGLE_API_KEY=your-api-key-here
        ```
        Get your API key from: https://makersuite.google.com/app/apikey
        
        **Option 2: Use Vertex AI**
        ```env
        GOOGLE_GENAI_USE_VERTEXAI=TRUE
        GOOGLE_CLOUD_PROJECT=your-project-id
        GOOGLE_CLOUD_LOCATION=us-central1
        ```
        """)
        return
    
    # Check system readiness
    if not ADK_AVAILABLE:
        st.error("System is not properly configured. Please ensure all required packages are installed.")
        st.info("Required packages: google-adk, pypdf, python-docx")
        return

    # File upload section in sidebar
    with st.sidebar:
        st.header("📄 Upload Requirements")
        
        # Clear any existing file uploader state on page refresh
        if 'file_uploader_key' not in st.session_state:
            st.session_state.file_uploader_key = 0
        
        uploaded_file = st.file_uploader(
            "Choose a .txt, .pdf, or .docx file",
            type=["txt", "pdf", "docx"],
            help="Upload your requirements document to begin analysis",
            key=f"file_uploader_{st.session_state.file_uploader_key}"
        )

        if uploaded_file is not None:
            st.success(f"Uploaded: {uploaded_file.name}")
            
            # Display file info
            file_size = len(uploaded_file.getvalue())
            st.info(f"File size: {file_size:,} bytes")
            
            print("Session_id before call:", st.session_state[ADK_SESSION_KEY])
            # Check file size limit (10MB)
            max_size = 10 * 1024 * 1024  # 10MB
            if file_size > max_size:
                st.error(f"File too large. Maximum size is {max_size // (1024*1024)}MB")
            else:
                if st.button("🚀 Start Analysis", type="primary", key="start_analysis"):
                    reset_run_state()
                    handle_context_retrieval(uploaded_file)
                    st.rerun()
        
        # RAG Compliance Documents Section
        st.divider()
        st.header("📚 Compliance Documents (RAG)")
        
        # RAG Stats
        try:
            rag_service = get_rag_service()
            if rag_service:
                stats = rag_service.get_collection_stats()
                st.info(f"📊 {stats['total_documents']} documents loaded\n📄 {stats['total_chunks']} document chunks")
                
                if stats['documents']:
                    with st.expander("View loaded documents"):
                        for doc in stats['documents']:
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                st.text(f"• {doc}")
                            with col2:
                                if st.button("🗑️", key=f"delete_{doc}", help=f"Delete {doc}"):
                                    if rag_service.delete_document(doc):
                                        st.success(f"Deleted {doc}")
                                        st.rerun()
            else:
                st.warning("⚠️ RAG service not available. Install: pip install chromadb sentence-transformers")
        except Exception as e:
            logger.warning(f"RAG service error: {e}")
            st.warning(f"⚠️ RAG service error: {str(e)[:100]}")
        
        # Upload compliance documents
        compliance_file = st.file_uploader(
            "Upload Compliance Documents",
            type=["txt", "pdf", "docx"],
            help="Upload FDA, ISO, IEC, HIPAA or other compliance documents for RAG",
            key="compliance_uploader"
        )
        
        if compliance_file is not None:
            file_size = len(compliance_file.getvalue())
            st.info(f"File: {compliance_file.name} ({file_size:,} bytes)")
            if st.button("📥 Ingest into RAG", key="ingest_compliance"):
                if handle_compliance_document_upload(compliance_file):
                    st.rerun()

    # Main application flow
    if st.session_state.processing_complete:
        display_results()
    elif st.session_state.context:
        st.header("📝 Review and Edit Requirements & Context")
        
        requirements = st.session_state.context.get('requirements', [])
        context_text = st.session_state.context.get('context', '')
        saved_requirement_contexts = st.session_state.context.get('requirement_contexts', {})
        
        if requirements:
            st.subheader("📋 Extracted Requirements (Editable)")
            st.info(f"Found {len(requirements)} requirements")
            
            # Create editable requirements list
            edited_requirements = []
            requirement_contexts = {}
            
            # Initialize requirement contexts in session state if not exists
            if 'requirement_contexts' not in st.session_state:
                st.session_state.requirement_contexts = {}
            
            # Restore saved requirement contexts if available
            if saved_requirement_contexts:
                for req_id, req_ctx in saved_requirement_contexts.items():
                    req_context_key = f"req_context_{req_id}"
                    if req_context_key not in st.session_state.requirement_contexts:
                        st.session_state.requirement_contexts[req_context_key] = req_ctx
            
            for i, req in enumerate(requirements):
                with st.expander(f"Requirement {i+1}: {req.get('id', f'REQ-{i+1:03d}')}", expanded=False):
                    # Editable Requirement ID
                    req_id = st.text_input(
                        f"Requirement ID {i+1}",
                        value=req.get('id', f'REQ-{i+1:03d}'),
                        key=f"req_id_{i}",
                        help="Edit the requirement ID"
                    )
                    
                    # Editable Requirement Text
                    req_text = st.text_area(
                        f"Requirement Text {i+1}",
                        value=req.get('text', 'No text available'),
                        height=150,
                        key=f"req_text_{i}",
                        help="Edit the requirement text"
                    )
                    
                    # Context for this specific requirement
                    # Try to get saved context for this requirement ID, fallback to empty string
                    original_req_id = req.get('id', f'REQ-{i+1:03d}')
                    req_context_key = f"req_context_{original_req_id}"
                    
                    # Check if we have saved context for this requirement
                    saved_context = saved_requirement_contexts.get(original_req_id, "")
                    if req_context_key not in st.session_state.requirement_contexts:
                        st.session_state.requirement_contexts[req_context_key] = saved_context
                    
                    # Use session state value if available, otherwise use saved context
                    current_context = st.session_state.requirement_contexts.get(req_context_key, saved_context)
                    
                    req_context = st.text_area(
                        f"Context for {req_id}",
                        value=current_context,
                        height=200,
                        key=f"req_context_input_{i}",
                        help="Edit the context associated with this requirement"
                    )
                    
                    # Update session state for this requirement's context (using new ID if changed)
                    new_req_context_key = f"req_context_{req_id}"
                    st.session_state.requirement_contexts[new_req_context_key] = req_context
                    
                    # Store edited requirement
                    edited_requirements.append({
                        'id': req_id,
                        'text': req_text
                    })
                    
                    requirement_contexts[req_id] = req_context
            
            st.divider()
            
            # Global Context Section
            st.subheader("🧠 Global Context")
            edited_context = st.text_area(
                "Review and edit the overall context gathered by the agent.",
                value=context_text,
                height=300,
                help="This context will be used by other agents for compliance and analysis"
            )
            
            # Save Button
            col1, col2, col3 = st.columns([1, 1, 2])
            with col1:
                if st.button("💾 Save Changes", type="primary", key="save_context"):
                    # Update session state with edited data
                    updated_context_data = {
                        'requirements': edited_requirements,
                        'context': edited_context,
                        'requirement_contexts': requirement_contexts  # Store per-requirement contexts
                    }
                    st.session_state.context = updated_context_data
                    
                    # Save to database
                    if db:
                        try:
                            db.store_session_data(st.session_state[ADK_SESSION_KEY], updated_context_data)
                            logger.info("Context data updated and saved to database")
                            st.success("✅ Changes saved successfully to database!")
                        except Exception as e:
                            logger.warning(f"Failed to save to database: {e}")
                            st.error(f"Failed to save to database: {e}")
                    else:
                        st.warning("⚠️ Database not available - changes saved to session only")
                        st.success("✅ Changes saved to session state!")
            
            with col2:
                if st.button("✅ Approve & Generate Test Cases", type="primary"):
                    if not edited_context.strip():
                        st.warning("Please provide context before proceeding.")
                    else:
                        # Update session state before generating
                        updated_context_data = {
                            'requirements': edited_requirements,
                            'context': edited_context,
                            'requirement_contexts': requirement_contexts
                        }
                        st.session_state.context = updated_context_data
                        
                        # Save to database before proceeding
                        if db:
                            try:
                                db.store_session_data(st.session_state[ADK_SESSION_KEY], updated_context_data)
                                logger.info("Context data saved before test case generation")
                            except Exception as e:
                                logger.warning(f"Failed to save to database: {e}")
                        
                        run_generation_pipeline(edited_context, updated_context_data)
                        st.rerun()
        else:
            st.warning("No requirements extracted from the document.")
            st.subheader("🧠 Generated Context")
            context_text = st.session_state.context.get('context', '')
            edited_context = st.text_area(
                "Review and edit the context gathered by the agent.",
                value=context_text,
                height=300,
                help="This context will be used by other agents for compliance and analysis"
            )
            
            # Save Button for context only
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("💾 Save Context", type="primary", key="save_context_only"):
                    updated_context_data = {
                        'requirements': [],
                        'context': edited_context
                    }
                    st.session_state.context = updated_context_data
                    
                    if db:
                        try:
                            db.store_session_data(st.session_state[ADK_SESSION_KEY], updated_context_data)
                            st.success("✅ Context saved successfully to database!")
                        except Exception as e:
                            st.error(f"Failed to save to database: {e}")
                    else:
                        st.success("✅ Context saved to session state!")
            
            with col2:
                if st.button("✅ Approve & Generate Test Cases", type="primary"):
                    if not edited_context.strip():
                        st.warning("Please provide context before proceeding.")
                    else:
                        # Create updated_context_data with empty requirements
                        updated_context_data = {
                            'requirements': [],
                            'context': edited_context
                        }
                        st.session_state.context = updated_context_data
                        
                        # Save to database before proceeding
                        if db:
                            try:
                                db.store_session_data(st.session_state[ADK_SESSION_KEY], updated_context_data)
                                logger.info("Context data saved before test case generation")
                            except Exception as e:
                                logger.warning(f"Failed to save to database: {e}")
                        
                        run_generation_pipeline(edited_context, updated_context_data)
                        st.rerun()
    else:
        st.info("Please upload a requirement document in the sidebar and click 'Start Analysis' to begin.")
        
        # Show sample format
        with st.expander("📄 Sample Document Format"):
            st.markdown("""
            Your document should contain clear requirements such as:
            
            - **Functional Requirements**: What the system should do
            - **Non-functional Requirements**: Performance, security, usability criteria  
            - **Business Rules**: Constraints and validation rules
            - **Compliance Requirements**: Regulatory standards to meet
            
            Example:
            ```
            REQ-001: The system shall authenticate users within 3 seconds
            REQ-002: Patient data must be encrypted using AES-256
            REQ-003: The system must comply with HIPAA regulations
            ```
            """)

def set_app_background(image_path):
    """
    Sets a background image for the Streamlit app and adjusts content background.
    """
    if not os.path.exists(image_path):
        logger.warning(f"Background image not found at: {image_path}")
        return

    try:
        with open(image_path, "rb") as f:
            img_bytes = f.read()
        encoded = base64.b64encode(img_bytes).decode()

        background_style = f"""
        <style>
        .stApp {{
            background-image: url('data:image/png;base64,{encoded}');
            background-size: cover;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }}
        /* Make content background semi-transparent */
        [data-testid="stAppViewContainer"] > .main {{
            background-color: rgba(255, 255, 255, 0.9);
        }}
        </style>
        """
        st.markdown(background_style, unsafe_allow_html=True)
    except Exception as e:
        logger.error(f"Failed to set background image: {e}")


if __name__ == "__main__":
    # Assuming your image is at 'images/background.png' relative to the app.py file.
    # Please change the filename if it's different.
    #set_app_background('src/images/img7.png')
    # st.markdown("""
    # <style>
    # .stApp {
    #     background-image: url('data:image/png;base64,%s');
    #     background-size: cover;
    #     background-position: center;
    # }
    # </style>
    # """ % base64.b64encode(open('src/images/img4.png', 'rb').read()).decode(), 
    # unsafe_allow_html=True)
    # Use an absolute path for reliability
    main()